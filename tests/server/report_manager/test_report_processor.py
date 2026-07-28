import base64
import re
import zipfile
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import markdown
import pytest

from openjiuwen_deepsearch.algorithm.report_export.conversion_utils import (
    normalize_docx_tables,
    postprocess_html,
    preprocess_markdown_text,
    protect_math_spans,
    restore_math_spans,
    wrap_html_tables,
)
from openjiuwen_deepsearch.algorithm.report_export.docx_export import convert_md_to_docx
from openjiuwen_deepsearch.algorithm.report_export.html_export import convert_md_to_html
from openjiuwen_deepsearch.algorithm.report_export.mermaid_preprocess import (
    MermaidRenderOptions,
    extract_xychart_metadata,
    preprocess_mermaid_code,
)
from openjiuwen_deepsearch.algorithm.report_export.report_bundle import build_report_bundle
from openjiuwen_deepsearch.algorithm.report_export.word_utils import (
    _normalize_latex_for_omml,
    html_to_doc,
    set_global_styles,
)


TINY_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO"
    "+/p9sAAAAASUVORK5CYII="
)

STYLE_MAP = {
    "heading1": "heading 1",
    "heading2": "heading 2",
    "heading3": "heading 3",
    "heading4": "heading 4",
    "heading5": "heading 5",
    "heading6": "heading 6",
    "heading7": "heading 7",
    "heading8": "heading 8",
    "heading9": "heading 9",
    "paragraph": "Normal",
    "table": "Table Grid",
    "default": "Normal",
}


def _convert_html_text(markdown_text: str, tmp_path: Path) -> str:
    """将 Markdown 转换为公共 legacy HTML 并返回文本。

    Args:
        markdown_text: 待转换 Markdown。
        tmp_path: pytest 临时工作目录。

    Returns:
        生成的 HTML 文本。
    """
    markdown_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    markdown_path.write_text(markdown_text, encoding="utf-8")
    convert_md_to_html(markdown_path, html_path)
    return html_path.read_text(encoding="utf-8")


def _convert_docx_document(markdown_text: str, tmp_path: Path) -> Document:
    """将 Markdown 转换为公共 DOCX 并返回文档对象。

    Args:
        markdown_text: 待转换 Markdown。
        tmp_path: pytest 临时工作目录。

    Returns:
        已生成的 Word 文档。
    """
    markdown_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    markdown_path.write_text(markdown_text, encoding="utf-8")
    convert_md_to_docx(markdown_path, docx_path)
    return Document(docx_path)


def _html_to_word(html_text: str) -> Document:
    """使用公共 HTML-to-DOCX 工具转换 HTML 文本。

    Args:
        html_text: 待转换 HTML。

    Returns:
        生成的 Word 文档。
    """
    document = Document()
    set_global_styles(document)
    html_to_doc(document, html_text, STYLE_MAP)
    return document


def test_set_global_styles_uses_compact_line_spacing():
    """Validate generated DOCX paragraphs use compact line spacing."""
    document = Document()

    set_global_styles(document)

    paragraph_format = document.styles["Normal"].paragraph_format
    assert paragraph_format.line_spacing_rule == WD_LINE_SPACING.MULTIPLE
    assert paragraph_format.line_spacing == 1.15


def test_preprocess_mermaid_code_scales_xychart_and_extracts_metadata():
    """Validate xychart preprocessing keeps parity with the reference offline flow.

    Returns:
        None.
    """
    processed, supplement = preprocess_mermaid_code(
        "xychart-beta\n  bar [1200]\n",
        MermaidRenderOptions(),
    )
    metadata = extract_xychart_metadata(processed, warn_on_invalid=False)

    assert supplement == ""
    assert 'y-axis "x1e3"' in processed
    assert "bar [1.2]" in processed
    assert len(metadata.series) == 1
    assert metadata.series[0].display_values == ["1.2"]


def test_preprocess_markdown_text_strips_internal_citation_markers():
    """Validate export preprocessing hides internal citation control markers.

    Returns:
        None.
    """
    text = (
        "保留引用[checked_citation:4][[5]](https://example.com/source)\n\n"
        "移除旧标记[citation:2]"
    )

    processed = preprocess_markdown_text(text)

    assert "checked_citation" not in processed
    assert "[citation:2]" not in processed
    assert '[5]</a>' in processed


def test_wrap_html_tables_adds_centering_container_once():
    """Validate HTML table wrapping is idempotent.

    Returns:
        None.
    """
    html_text = "<p>intro</p><table><tr><td>A</td></tr></table>"

    processed = wrap_html_tables(html_text)
    processed_twice = wrap_html_tables(processed)

    assert '<div class="table-wrap"><table>' in processed
    assert processed.count('class="table-wrap"') == 1
    assert processed_twice.count('class="table-wrap"') == 1


def test_postprocess_html_wraps_tables_without_rewriting_svg():
    """Validate table wrapping does not mutate Mermaid SVG HTML.

    Returns:
        None.
    """
    html_text = (
        '<div class="mermaid-rendered"><svg viewBox="0 0 100 100"></svg></div>'
        "<table><tr><td>A</td></tr></table>"
    )

    processed = postprocess_html(html_text)

    assert 'viewBox="0 0 100 100"' in processed
    assert '<div class="table-wrap"><table>' in processed


def test_report_html_convert_from_markdown_wraps_tables(tmp_path):
    """Validate direct HTML conversion wraps Markdown tables.

    Returns:
        None.
    """
    html_text = _convert_html_text("| A | B |\n|---|---|\n| 1 | 2 |", tmp_path)

    assert 'class="table-wrap"' in html_text
    assert "<table>" in html_text


def test_report_html_convert_from_markdown_uses_shared_safe_math_handling(tmp_path):
    """Validate direct HTML conversion shares offline math/currency behavior."""
    html_text = _convert_html_text("变量 $G$ 保留为公式，价格 $4 和 $5 保持文本。", tmp_path)

    assert "$G$" in html_text
    assert "$4 和 $5" in html_text
    assert "katex" in html_text


def test_report_word_convert_from_markdown_keeps_wrapped_tables(tmp_path):
    """Validate online DOCX conversion keeps tables wrapped for HTML centering.

    Returns:
        None.
    """
    doc = _convert_docx_document("| A | B |\n|---|---|\n| 1 | 2 |", tmp_path)

    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "A"
    assert doc.tables[0].cell(1, 1).text == "2"


def test_report_word_convert_from_markdown_uses_shared_safe_math_handling(tmp_path):
    """Validate direct DOCX conversion renders math without converting currency."""
    doc = _convert_docx_document("变量 $G$ 保留为公式，价格 $4 和 $5 保持文本。", tmp_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert r"\(G\)" not in paragraph_text
    assert "$4 和 $5" in paragraph_text


def test_report_word_convert_from_html_handles_irregular_table_rows():
    """Validate HTML table conversion tolerates rows with different cell counts."""
    html_text = (
        '<div class="report-container">'
        "<table>"
        "<tr><th>A</th><th>B</th></tr>"
        "<tr><td>1</td><td>2</td><td>3</td></tr>"
        "</table>"
        "</div>"
    )

    doc = _html_to_word(html_text)

    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 3
    assert doc.tables[0].cell(0, 2).text == ""
    assert doc.tables[0].cell(1, 2).text == "3"


def test_report_word_convert_from_html_limits_nested_block_depth():
    """Validate deeply nested HTML is flattened instead of recursing indefinitely."""
    html_text = (
        '<div class="report-container">'
        + "<div>" * 120
        + "<p>深层内容</p>"
        + "</div>" * 120
        + "</div>"
    )

    doc = _html_to_word(html_text)

    assert any("深层内容" in paragraph.text for paragraph in doc.paragraphs)


def test_html_to_doc_embeds_relative_images_from_base_path(tmp_path):
    """Validate pure-Python HTML conversion embeds bundle-local image files."""
    image_dir = tmp_path / "charts"
    image_dir.mkdir()
    (image_dir / "chart_0.png").write_bytes(base64.b64decode(TINY_PNG_BASE64))
    document = Document()
    style_map = {
        "heading1": "heading 1",
        "heading2": "heading 2",
        "heading3": "heading 3",
        "heading4": "heading 4",
        "heading5": "heading 5",
        "heading6": "heading 6",
        "heading7": "heading 7",
        "heading8": "heading 8",
        "heading9": "heading 9",
        "paragraph": "Normal",
        "table": "Table Grid",
        "default": "Normal",
    }

    html_to_doc(
        document,
        '<div class="report-container"><p><img src="charts/chart_0.png" alt="Chart"></p></div>',
        style_map,
        base_path=tmp_path,
    )

    assert len(document.inline_shapes) == 1


def test_html_to_doc_limits_large_images_to_page_width(tmp_path):
    """Validate pure-Python DOCX conversion scales oversized images to fit the page."""
    image = pytest.importorskip("PIL.Image")
    image_dir = tmp_path / "charts"
    image_dir.mkdir()
    image_path = image_dir / "large_chart.png"
    image.new("RGB", (2400, 800), color="white").save(image_path)
    document = Document()
    style_map = {
        "heading1": "heading 1",
        "heading2": "heading 2",
        "heading3": "heading 3",
        "heading4": "heading 4",
        "heading5": "heading 5",
        "heading6": "heading 6",
        "heading7": "heading 7",
        "heading8": "heading 8",
        "heading9": "heading 9",
        "paragraph": "Normal",
        "table": "Table Grid",
        "default": "Normal",
    }

    html_to_doc(
        document,
        '<div class="report-container"><p><img src="charts/large_chart.png" alt="Chart"></p></div>',
        style_map,
        base_path=tmp_path,
    )

    max_width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
    assert document.inline_shapes[0].width <= max_width


def test_convert_md_to_docx_embeds_relative_images(tmp_path):
    """Validate pure-Python DOCX export embeds bundle-local image files."""
    md_path = tmp_path / "report.md"
    image_dir = tmp_path / "charts"
    image_dir.mkdir()
    (image_dir / "chart_0.png").write_bytes(base64.b64decode(TINY_PNG_BASE64))
    docx_path = tmp_path / "report.docx"
    md_path.write_text("# Title\n\n![Chart](charts/chart_0.png)\n", encoding="utf-8")

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    assert docx_path.exists()
    assert len(document.inline_shapes) == 1


def test_convert_md_to_docx_keeps_citations_as_superscript_links(tmp_path):
    """Validate DOCX export preserves citation links as superscript text."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text("正文需要引用[[1]](https://example.com/source)。\n", encoding="utf-8")

    convert_md_to_docx(md_path, docx_path)

    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert '<w:vertAlign w:val="superscript"/>' in document_xml
    assert "[1]" in document_xml


def test_convert_md_to_docx_keeps_code_blocks(tmp_path):
    """Validate DOCX export preserves fenced code block content."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text("# Title\n\n```python\nprint('hello')\n```\n", encoding="utf-8")

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    assert any("print('hello')" in paragraph.text for paragraph in document.paragraphs)


def test_convert_md_to_docx_keeps_citation_links_in_math_paragraphs(tmp_path):
    """Validate math paragraphs keep inline citation links and superscript formatting."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text("公式 $x^2$ 后引用[[1]](https://example.com/source)。\n", encoding="utf-8")

    convert_md_to_docx(md_path, docx_path)

    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
        document_rels = zip_file.read("word/_rels/document.xml.rels").decode("utf-8")
    assert '<w:vertAlign w:val="superscript"/>' in document_xml
    assert "[1]" in document_xml
    assert "https://example.com/source" in document_rels


def test_convert_md_to_docx_keeps_table_cell_links(tmp_path):
    """Validate DOCX export preserves hyperlinks inside table cells."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "| Source |\n|---|\n| [link](https://example.com/source) |\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    with zipfile.ZipFile(docx_path) as zip_file:
        document_rels = zip_file.read("word/_rels/document.xml.rels").decode("utf-8")
    assert document.tables[0].cell(1, 0).text == "link"
    assert "https://example.com/source" in document_rels


def test_convert_md_to_html_keeps_indented_markdown_list_items_out_of_code_blocks(tmp_path):
    """Validate report-style indented bullet lines render as lists, not code blocks."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "\n".join(
            [
                "![Chart](charts/chart.png)",
                "<font size=2>**图表**: 说明</font>[[1]](https://example.com/source)",
                "    - **逻辑二：软硬解耦与异构全栈底座打破算力垄断**。构建兼容异构芯片的软件栈至关重要[[1]](https://example.com/source)。",
            ]
        ),
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "<pre><code>" not in html_text
    assert "<li>" in html_text
    assert "逻辑二：软硬解耦" in html_text


def test_convert_md_to_html_keeps_consecutive_indented_list_items_at_same_level(tmp_path):
    """Validate consecutive report-style indented bullets stay sibling list items."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "\n".join(
            [
                "Caption line",
                "    - first item",
                "    - second item",
            ]
        ),
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "<pre><code>" not in html_text
    assert html_text.count("<ul>") == 1
    assert html_text.count("<li>") == 2
    assert "first item" in html_text
    assert "second item" in html_text


def test_html_to_doc_keeps_nested_list_items_as_separate_paragraphs():
    """Validate nested HTML lists do not collapse child item text into the parent paragraph."""
    document = Document()
    style_map = {
        "heading1": "heading 1",
        "heading2": "heading 2",
        "heading3": "heading 3",
        "heading4": "heading 4",
        "heading5": "heading 5",
        "heading6": "heading 6",
        "heading7": "heading 7",
        "heading8": "heading 8",
        "heading9": "heading 9",
        "paragraph": "Normal",
        "table": "Table Grid",
        "default": "Normal",
    }

    html_to_doc(
        document,
        (
            '<div class="report-container">'
            "<ul><li>first<ul><li>second</li></ul></li><li>third</li></ul>"
            "</div>"
        ),
        style_map,
    )

    assert [paragraph.text for paragraph in document.paragraphs] == ["first", "second", "third"]


def test_html_to_doc_indents_nested_list_items():
    """Validate DOCX paragraphs preserve nested list depth and native bullets."""
    document = Document()
    style_map = {
        "paragraph": "Normal",
        "default": "Normal",
    }

    html_to_doc(
        document,
        (
            '<div class="report-container">'
            "<ul><li>parent<ul><li>child</li></ul></li><li>sibling</li></ul>"
            "</div>"
        ),
        style_map,
    )

    parent, child, sibling = document.paragraphs
    assert parent.paragraph_format.left_indent is None
    assert child.paragraph_format.left_indent.pt == 18
    assert sibling.paragraph_format.left_indent is None
    parent_num_pr = parent._p.pPr.numPr
    child_num_pr = child._p.pPr.numPr
    sibling_num_pr = sibling._p.pPr.numPr
    assert parent_num_pr.numId.val == child_num_pr.numId.val == sibling_num_pr.numId.val
    assert parent_num_pr.ilvl.val == 0
    assert child_num_pr.ilvl.val == 1
    assert sibling_num_pr.ilvl.val == 0
    numbering_xml = document.part.numbering_part.element.xml
    assert 'w:numFmt w:val="bullet"' in numbering_xml
    assert 'w:lvlText w:val="•"' in numbering_xml
    assert 'w:lvlText w:val="◦"' in numbering_xml


def test_html_to_doc_uses_native_numbering_for_ordered_lists():
    """Validate ordered HTML lists become native multilevel Word numbering."""
    document = Document()

    html_to_doc(
        document,
        '<div class="report-container"><ol><li>first<ol><li>child</li></ol></li><li>second</li></ol></div>',
        {"paragraph": "Normal", "default": "Normal"},
    )

    first, child, second = document.paragraphs
    first_num_pr = first._p.pPr.numPr
    child_num_pr = child._p.pPr.numPr
    second_num_pr = second._p.pPr.numPr
    assert first_num_pr.numId.val == child_num_pr.numId.val == second_num_pr.numId.val
    assert first_num_pr.ilvl.val == 0
    assert child_num_pr.ilvl.val == 1
    assert second_num_pr.ilvl.val == 0
    numbering_xml = document.part.numbering_part.element.xml
    assert 'w:numFmt w:val="decimal"' in numbering_xml
    assert 'w:lvlText w:val="%1."' in numbering_xml


def test_report_table_css_preserves_global_width_and_centers_wrapped_tables():
    """Validate report CSS limits table centering changes to wrapped tables.

    Returns:
        None.
    """
    from openjiuwen_deepsearch.algorithm.report_export.html_export import HTML_TEMPLATE

    css_text = HTML_TEMPLATE

    assert re.search(r"table\s*\{[^}]*width:\s*fit-content;", css_text, flags=re.DOTALL)
    assert re.search(
        r"\.table-wrap\s+table\s*\{[^}]*width:\s*auto;[^}]*max-width:\s*100%;[^}]*margin:\s*0\s+auto;",
        css_text,
        flags=re.DOTALL,
    )


def test_report_html_export_renders_mermaid_or_falls_back(tmp_path):
    """Validate HTML export renders Mermaid or preserves source as fallback.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    final_result = {
        "response_content": "# 标题\n\n```mermaid\ngraph TD\nA-->B\n```",
        "infer_messages": [],
        "chart_messages": [],
        "warning_info": "",
        "exception_info": "",
    }

    bundle = build_report_bundle(final_result, tmp_path)
    html_path = bundle.root_dir / "report.html"
    convert_md_to_html(bundle.markdown_path, html_path)
    html_text = html_path.read_text(encoding="utf-8")

    assert "<html" in html_text.lower()
    assert "标题" in html_text
    assert ("<svg" in html_text) or ("language-mermaid" in html_text)


def test_normalize_docx_tables_centers_tables_and_captions(tmp_path):
    """Validate DOCX table normalization centers table objects and captions.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    docx_path = tmp_path / "tables.docx"
    document = Document()
    document.add_paragraph("普通正文段落")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    document.add_paragraph("表2-1：合肥市“三电”系统核心企业的技术实力与市场表现")
    document.save(docx_path)

    normalize_docx_tables(docx_path)

    normalized = Document(docx_path)
    assert normalized.tables[0].alignment == WD_TABLE_ALIGNMENT.CENTER
    assert normalized.paragraphs[0].alignment is None
    assert normalized.paragraphs[-1].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_convert_md_to_html_annotates_xychart_value_labels(tmp_path):
    """Validate HTML export annotates xychart SVG output with value labels.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "```mermaid\n---\nconfig:\n    showDataLabel: true\n---\n"
        "xychart-beta\n"
        '  x-axis ["收入"]\n'
        '  y-axis "亿元" 0 --> 2\n'
        "  bar [1.2]\n```",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "chart-value-label" in html_text


def test_convert_md_to_html_keeps_legacy_font_caption_separate_from_following_list(tmp_path):
    """Validate legacy font captions do not absorb following bullet lists."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "![图表示例](chart.png)\n"
        "<font size=2>**图表示例**: 图注说明</font>[[1]](https://example.com)\n"
        "- **技术维度**：第一条\n"
        "- **经济维度**：第二条\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert ".figure-caption {" in html_text
    assert "width: 100%;" in html_text
    assert "text-align: center;" in html_text
    assert '<div class="figure-caption">' in html_text
    assert "<ul>" in html_text
    assert "<li><strong>技术维度</strong>" in html_text


def test_convert_md_to_html_separates_paragraph_from_following_bullets_without_blank_line(tmp_path):
    """Validate list items render after a paragraph even when the source misses a blank line."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "在代际价值观层面，这一人群展现出强烈的矛盾统一体特征：\n"
        "- **求稳与求变并存**：第一条\n"
        "- **务实与悦己交织**：第二条\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "<p>在代际价值观层面，这一人群展现出强烈的矛盾统一体特征：</p>" in html_text
    assert "<ul>" in html_text
    assert "<li><strong>求稳与求变并存</strong>：第一条</li>" in html_text


def test_convert_md_to_html_separates_paragraph_from_following_table_without_blank_line(tmp_path):
    """Validate pipe tables render after a paragraph even when the source misses a blank line."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "表2-1梳理了测试数据：\n"
        "| 列1 | 列2 |\n"
        "| --- | --- |\n"
        "| A | B |\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "<p>表2-1梳理了测试数据：</p>" in html_text
    assert '<div class="table-wrap"><table>' in html_text
    assert "| 列1 | 列2 |" not in html_text


def test_convert_md_to_html_centers_table_display(tmp_path):
    """Validate exported HTML uses centered table presentation styles."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "| 列1 | 列2 |\n"
        "| --- | --- |\n"
        "| A | B |\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "table {" in html_text
    assert "margin: 16px auto 24px;" in html_text
    assert "width: fit-content;" in html_text
    assert "max-width: 100%;" in html_text
    assert "th[style], td[style] {" in html_text
    assert "text-align: center !important;" in html_text
    assert "text-align: center;" in html_text


def test_report_docx_export_creates_docx_file(tmp_path):
    """Validate DOCX export writes a pure-Python generated file into the bundle.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    final_result = {
        "response_content": "# Title\n\nPlain text.",
        "infer_messages": [],
        "chart_messages": [],
        "warning_info": "",
        "exception_info": "",
    }

    bundle = build_report_bundle(final_result, tmp_path)
    docx_path = bundle.root_dir / "report.docx"
    convert_md_to_docx(bundle.markdown_path, docx_path)

    assert docx_path.exists()
    document = Document(docx_path)
    assert any(paragraph.text == "Title" for paragraph in document.paragraphs)
    assert any(paragraph.text == "Plain text." for paragraph in document.paragraphs)


def test_convert_md_to_docx_normalizes_headings_fonts_and_tables(tmp_path, monkeypatch):
    """Validate DOCX export uses heading/font/table post-processing flow.

    Args:
        tmp_path: pytest 提供的临时目录。
        monkeypatch: pytest monkeypatch fixture。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text("1. Heading\n", encoding="utf-8")

    font_calls = {"count": 0}
    table_calls = {"count": 0}

    monkeypatch.setattr(
        "openjiuwen_deepsearch.algorithm.report_export.docx_export.normalize_docx_fonts",
        lambda *_args, **_kwargs: font_calls.__setitem__("count", font_calls["count"] + 1),
        raising=False,
    )
    monkeypatch.setattr(
        "openjiuwen_deepsearch.algorithm.report_export.docx_export.normalize_docx_tables",
        lambda *_args, **_kwargs: table_calls.__setitem__("count", table_calls["count"] + 1),
        raising=False,
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    assert any(paragraph.text == "1 Heading" for paragraph in document.paragraphs)
    assert font_calls["count"] == 1
    assert table_calls["count"] == 1


def test_convert_md_to_docx_preserves_short_bold_spans_in_long_chinese_summary(tmp_path):
    """Validate DOCX export keeps inline bold spans in long Chinese summary paragraphs."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# 摘要\n\n"
        "都市年轻职场人群（22-35岁）在职业周期演进与高居住成本制约下呈现显著分化，"
        "其核心生存图景由**69.6%**企业新招毕业生硕士占比更高、一线城市高达**45.6%**"
        "的房租负担率及**79.2%**企业离职率低于**10%**的求稳心态共同刻画；该群体深陷"
        "工作高压（**31.5%**日工作超10小时）、时间剥夺（北京单程通勤**47分钟**）、社交"
        "萎缩（超三成频繁孤独）与生活品质坍塌（超**90%**受亚健康影响）交织的痛点因果网，"
        "驱动其行为向情绪补剂常态化（近九成为情绪买单）、社交模块化（**54.4%**有搭子）"
        "与技术双刃剑（超**56%**日常使用GenAI但超六成担忧失业）三大策略演化；由此催生"
        "效率工具（**56.1%**愿为AI付费）、情绪价值（四成向AI倾诉）、零糖社交与零家务闭环"
        "等复合需求，其消费决策呈现极致折叠的精算师特质（比价工具使用率达**78%**）与为情绪"
        "溢价买单并存，复购核心由体验确证（**77.6%**因使用感好复购）与情绪持续供给驱动。\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    summary_paragraph = document.paragraphs[1]
    bold_runs = {run.text for run in summary_paragraph.runs if run.text and run.bold}

    assert "**" not in summary_paragraph.text
    assert {
        "69.6%",
        "45.6%",
        "79.2%",
        "10%",
        "31.5%",
        "47分钟",
        "90%",
        "54.4%",
        "56%",
        "56.1%",
        "78%",
        "77.6%",
    }.issubset(bold_runs)


def test_convert_md_to_docx_raises_file_not_found_for_missing_markdown(tmp_path):
    """Validate DOCX export still surfaces missing Markdown input."""
    with pytest.raises(FileNotFoundError):
        convert_md_to_docx(tmp_path / "missing.md", tmp_path / "report.docx")


def test_convert_md_to_docx_fallback_on_unconvertible_latex(tmp_path):
    """Validate DOCX export does not crash when a LaTeX formula cannot be converted.

    Instead, the raw formula text should be preserved in the output.
    This covers cases like deeply nested unbalanced left/right delimiters
    that latex2mathml rejects with ExtraLeftOrMissingRightError.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    # Nested unbalanced \left...\right — latex2mathml raises ExtraLeftOrMissingRightError
    # Use raw string to avoid \t/\f etc. being interpreted as escape chars
    md_path.write_text(
        r"# 测试" + "\n\n"
        r"$$\left[\left(x\right)\right) - y\right)$$" + "\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    # Should contain the raw formula text as fallback (with $$ delimiters preserved)
    assert any(r"\left" in paragraph.text for paragraph in document.paragraphs)


def test_convert_md_to_docx_keeps_valid_math_inline_and_block(tmp_path):
    """Validate DOCX export keeps well-formed LaTeX formulas converted to OMML.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# 测试\n\n"
        "行内公式 $x^2 + y^2 = z^2$。\n\n"
        "$$E = mc^2$$\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    # OMML namespace should be present for converted formulas
    assert "http://schemas.openxmlformats.org/officeDocument/2006/math" in document_xml


def test_convert_md_to_docx_renders_binomial_power_display_math(tmp_path):
    """Validate DOCX export converts binomial expressions with outer powers."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# Test\n\n"
        "$$\n"
        r"\sum_{k=1}^{\infty} \frac{21k-8}{k^3 \binom{2k}{k}^3} = \frac{\pi^2}{6}"
        "\n$$\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "$$" not in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml


def test_convert_md_to_docx_renders_binomial_power_inline_math(tmp_path):
    """Validate inline DOCX math also normalizes binomial expressions."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# Test\n\n"
        r"The inline choice formula is $\binom{n}{r}^2$ in text."
        "\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert r"$\binom{n}{r}^2$" not in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml


def test_convert_md_to_docx_renders_aligned_display_math(tmp_path):
    """Validate DOCX export converts aligned display formulas."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# Test\n\n"
        "$$\n"
        r"\begin{aligned} "
        r"\frac {\mathrm {d} x}{\mathrm {d} t}&=\sigma y-\sigma x,\\ "
        r"\frac {\mathrm {d} y}{\mathrm {d} t}&=\rho x-xz-y,\\ "
        r"\frac {\mathrm {d} z}{\mathrm {d} t}&=xy-\beta z. "
        r"\end{aligned}"
        "\n$$\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "$$" not in paragraph_text
    assert r"\begin{aligned}" not in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml


def test_html_export_contains_katex_script(tmp_path):
    """Validate HTML export includes KaTeX script for formula rendering.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "# 测试\n\n"
        "公式：$$E = mc^2$$\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "katex" in html_text.lower()
    assert "auto-render" in html_text
    assert "renderMathInElement" in html_text


def test_protect_math_spans_preserves_underscores_through_markdown():
    """Validate math placeholders survive Markdown conversion without emphasis injection.

    Underscores and asterisks inside ``$...$`` / ``$$...$$`` must not be turned
    into ``<em>``/``<strong>`` tags by Python-Markdown.

    Returns:
        None.
    """
    text = (
        "正文 $$\\mathcal{J}_{GRPO}(\\theta) = \\sum_{i=1}^G \\pi_\\theta$$ "
        "行内 $a_{i} * b_{j}$ 结束"
    )

    protected, formulas = protect_math_spans(text)
    assert len(formulas) == 2
    assert "$" not in protected  # 公式分隔符已被占位符取代

    rendered = markdown.markdown(protected, extensions=["extra"], output_format="html5")
    restored = restore_math_spans(rendered, formulas)

    assert "$$\\mathcal{J}_{GRPO}(\\theta) = \\sum_{i=1}^G \\pi_\\theta$$" in restored
    assert "$a_{i} * b_{j}$" in restored
    assert "<em>" not in restored
    assert "<strong>" not in restored


def test_protect_math_spans_ignores_code_spans_and_fenced_blocks():
    """Validate math placeholders do not rewrite code examples."""
    text = (
        "正文 $x_i$。\n\n"
        "`price = \"$5\"`\n\n"
        "```python\n"
        "formula = \"$x_i$\"\n"
        "```\n"
    )

    protected, formulas = protect_math_spans(text)

    assert formulas == ["$x_i$"]
    assert '`price = "$5"`' in protected
    assert 'formula = "$x_i$"' in protected


def test_protect_math_spans_skips_currency_but_keeps_simple_variables():
    """Validate inline math protection distinguishes currency from variables."""
    text = (
        r"成本为 $5 到 $10，收入约 $1,200.50，区间 amount $1,200 to $1,300，"
        r"变量 $G$、$H0$、$xyz$、$a, b, m$、$a, b, c, ...$、$(x,y,z)$ 和公式 "
        r"$G = (V, E)$、$P(1, 2, 1, (1, 0))$、$1/2$、$4*3$、$n!$、$5!$、"
        r"$|z|$、$|x-y|$、$f'(t)$、$f''(t)$、$f_i'(t)$、"
        r"$F(s) G(s)$、$F_0(s)\,G_1(s)$、$||x||$、$\|x\|$、$\lVert x \rVert$ 需要保留，"
        r"百分比 $79.29\%$、$79.29%$ 也需要保留，"
        r"转义美元 \$8 不处理。"
    )

    protected, formulas = protect_math_spans(text)

    assert formulas == [
        "$G$",
        "$H0$",
        "$xyz$",
        "$a, b, m$",
        "$a, b, c, ...$",
        "$(x,y,z)$",
        "$G = (V, E)$",
        "$P(1, 2, 1, (1, 0))$",
        "$1/2$",
        "$4*3$",
        "$n!$",
        "$5!$",
        "$|z|$",
        "$|x-y|$",
        "$f'(t)$",
        "$f''(t)$",
        "$f_i'(t)$",
        "$F(s) G(s)$",
        "$F_0(s)\\,G_1(s)$",
        "$||x||$",
        "$\\|x\\|$",
        "$\\lVert x \\rVert$",
        "$79.29\\%$",
        "$79.29%$",
    ]
    assert "$5 到 $10" in protected
    assert "$1,200.50" in protected
    assert "$1,200 to $1,300" in protected
    assert r"\$8" in protected


def test_html_export_uses_safe_katex_delimiters_for_currency_text(tmp_path):
    """Validate KaTeX does not reinterpret currency dollars as formulas.

    ``$G$`` should be treated as math and wrapped as ``$G$`` in the output,
    while ``$4`` and ``$5`` are currency and should remain as-is.
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "# 测试\n\n"
        "其中，$G$为组内采样数量，每一组的价格分别为$4和$5。\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "$G$" in html_text
    assert "$4和$5" in html_text
    assert "\\\\(G\\\\)" not in html_text  # 不再是 MathJax 的 \(...\) 格式


def test_convert_md_to_html_keeps_dollar_unit_table_intact(tmp_path):
    """Validate dollar unit labels do not break Markdown table parsing."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "| Provider | Price ($/GPU-hr) | Note |\n"
        "| :--- | :--- | :--- |\n"
        "| Thunder Compute | $1.38 | fixed on-demand price |\n"
        "\n"
        "Later text before a parameter mention.\n"
        "\n"
        "Hybrid sampling parameter ($q_r$) should render as math.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html_text, "html.parser")
    assert soup.find("table") is not None
    assert "Price ($/GPU-hr)" in soup.get_text()
    assert "$q_r$" in html_text
    assert r"\\(/GPU-hr" not in html_text


def test_convert_md_to_html_protects_numeric_indicator_math_and_currency(tmp_path):
    """Validate numeric indicator formulas are not mistaken for currency."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "| Term | Formula | Price |\n"
        "| :--- | :--- | :--- |\n"
        "| torque | $1(\\tau_t \\notin [\\tau_{min}, \\tau_{max}])$ | $1.38 |\n"
        "\n"
        "Each group costs $4 and $5.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "$1(\\tau_t \\notin [\\tau_{min}, \\tau_{max}])$" in html_text
    assert "$1.38" in html_text
    assert "$4 and $5" in html_text


def test_convert_md_to_html_protects_comparison_math_before_markdown(tmp_path):
    """Validate comparison operators inside formulas do not skip math protection."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "| Term | Formula | Price |\n"
        "| :--- | :--- | :--- |\n"
        "| fall | $1(F^{left}_{feet}, F^{right}_{feet} < 1)$ | $1.38 |\n"
        "\n"
        "Scale uses $r_{t,i} < 0$ while escaped price \\$8 remains text.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    soup_text = BeautifulSoup(html_text, "html.parser").get_text(" ", strip=True)
    assert "$1(F^{left}_{feet}, F^{right}_{feet} &lt; 1)$" in html_text
    assert "$r_{t,i} &lt; 0$" in html_text
    assert "$1(F^{left}_{feet}, F^{right}_{feet} < 1)$" in soup_text
    assert "$r_{t,i} < 0$" in soup_text
    assert "<em" not in html_text
    assert "$1.38" in html_text
    assert r"\$8" in html_text


def test_convert_md_to_html_protects_numeric_scientific_math(tmp_path):
    """Validate numeric-leading scientific formulas are not treated as currency."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "Training converges at $4 \\times 10^4$ samples. "
        "Plain prices $4 and $5 remain text.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "$4 \\times 10^4$" in html_text
    assert "$4 and $5" in html_text


def test_convert_md_to_html_protects_numeric_arithmetic_math(tmp_path):
    """Validate numeric arithmetic formulas are not treated as currency."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "Arithmetic formulas $2+2=4$, $4-3=1$, and $4−3=1$ render as math. "
        "Plain prices $4 and $5 remain text.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "$2+2=4$" in html_text
    assert "$4-3=1$" in html_text
    assert "$4−3=1$" in html_text
    assert "$4 and $5" in html_text


def test_convert_md_to_html_escapes_formula_content_before_restore(tmp_path):
    """Validate formula text containing '<' does not become HTML tags."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "Compare $x<y$ and keep price $4 and $5 as text.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    soup_text = BeautifulSoup(html_text, "html.parser").get_text(" ", strip=True)
    assert "$x&lt;y$" in html_text
    assert "$x<y$" in soup_text
    assert "$4 and $5" in soup_text


def test_convert_md_to_docx_renders_numeric_arithmetic_math(tmp_path):
    """Validate DOCX export converts numeric arithmetic formulas to OMML."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "Arithmetic formulas $2+2=4$, $4-3=1$, and $4−3=1$ render as math. "
        "Plain prices $4 and $5 remain text.\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "$2+2=4$" not in paragraph_text
    assert "$4-3=1$" not in paragraph_text
    assert "$4−3=1$" not in paragraph_text
    assert "$4 and $5" in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert document_xml.count("<m:oMath") >= 3


def test_convert_md_to_docx_preserves_formula_text_with_less_than(tmp_path):
    """Validate DOCX export does not lose formula content after '<'."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "Compare $x<y$ and keep price $4 and $5 as text.\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "$x<y$" not in paragraph_text
    assert "Compare" in paragraph_text
    assert "$4 and $5" in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml
    assert "<m:t>x</m:t>" in document_xml
    assert "<m:t>y</m:t>" in document_xml


def test_convert_md_to_html_protects_equation_references(tmp_path):
    """Validate equation references wrapped in dollars render as inline math."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "The constraint follows $(Eq. 9)$ and the loss follows $(Equation 10)$.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "$(Eq. 9)$" in html_text
    assert "$(Equation 10)$" in html_text


def test_convert_md_to_html_protects_numeric_tuples_and_prime_variables(tmp_path):
    """Validate numeric tuples and prime variables render as inline math."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "Clamp range $(0.75, 1.5)$ and transformed variable $z'$ are formulas. "
        "Quoted text $foo'$ remains text.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "$(0.75, 1.5)$" in html_text
    assert "$z'$" in html_text
    assert "$foo'" in html_text


def test_convert_md_to_html_keeps_function_call_math_distinct_from_closing_dollars(tmp_path):
    """Validate function-call formulas do not leave broken dollar delimiters."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "Advantage Function, $A(s,a)$, Value Network, $V(s)$, "
        "and $A(s,a) = Q(s,a) - V(s)$ are formulas. Price $1.38 remains text.\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "$A(s,a)$" in html_text
    assert "$V(s)$" in html_text
    assert "$A(s,a) = Q(s,a) - V(s)$" in html_text
    assert "$1.38" in html_text


def test_html_export_defines_bm_macro(tmp_path):
    """Validate KaTeX can render reports that use the common \\bm command."""
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        r"Policy $\pi_{\bm{\theta}}:\mathcal{O}\rightarrow\mathcal{A}$.",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert r"$\pi_{\bm{\theta}}:\mathcal{O}\rightarrow\mathcal{A}$" in html_text
    # 宏键和宏值在最终 HTML/JS 源码中必须各含两个反斜杠，
    # 使 JavaScript 解析后为单反斜杠 \bm / \boldsymbol{#1}（KaTeX 宏键需带前导反斜杠）。
    assert r"'\\bm'" in html_text
    assert r"'\\boldsymbol{#1}'" in html_text


def test_convert_md_to_docx_preserves_currency_dollars_next_to_inline_math(tmp_path):
    """Validate DOCX conversion does not convert currency dollars as math."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# 测试\n\n"
        "其中，$G$为组内采样数量，每一组的价格分别为$4和$5。\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "$4和$5" in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "http://schemas.openxmlformats.org/officeDocument/2006/math" in document_xml
    assert r"\(G\)" not in document_xml


def test_convert_md_to_docx_renders_inline_math_without_currency_text(tmp_path):
    """Validate DOCX conversion renders normalized inline math without literal slashes."""
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text("# Test\n\nVariable $G$ should render as math.\n", encoding="utf-8")

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert r"\(G\)" not in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml


def test_html_export_preserves_underscore_math(tmp_path):
    """Validate HTML export keeps underscore-heavy formulas intact for KaTeX.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "# 测试\n\n"
        "$$\\sum_{i=1}^G \\pi_\\theta(o_i|q)$$\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    html_text = html_path.read_text(encoding="utf-8")
    assert "\\sum_{i=1}^G \\pi_\\theta(o_i|q)" in html_text
    # 公式未被 Markdown 拆成斜体标签
    assert "<em>" not in html_text
    assert "<strong>" not in html_text


def test_docx_export_preserves_underscore_math(tmp_path):
    """Validate DOCX export handles underscore-heavy formulas without corruption.

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# 测试\n\n"
        "$$x_{i}^{2} + y_{j}^{2} = z_{k}^{2}$$\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    # 公式应转为 OMML，且未残留 Markdown 注入的斜体标签
    assert "http://schemas.openxmlformats.org/officeDocument/2006/math" in document_xml


def test_is_likely_inline_math_recognizes_long_alphanumeric_variable():
    """长字母变量（如 velocityvelocityvelocity）应被识别为公式而非普通文本。

    回归测试：commit 867d0e6 引入的 `^[A-Za-zΑ-ω][A-Za-zΑ-ω0-9]{0,4}$` 限制
    字母变量最多 5 字符，导致 24 字符的 ``velocityvelocityvelocity`` 不被识别为
    公式，DOCX 中残留字面 ``$...$``。
    """
    from openjiuwen_deepsearch.algorithm.report_export.conversion_utils import (
        _is_likely_inline_math,
    )

    assert _is_likely_inline_math("velocityvelocityvelocity") is True
    assert _is_likely_inline_math("abcdefghijklmnopqrstuvwxyz") is True
    # 短变量仍应被识别
    assert _is_likely_inline_math("x") is True
    assert _is_likely_inline_math("xyz") is True
    # 含空格的普通文本不应被识别
    assert _is_likely_inline_math("plain text") is False


def test_convert_md_to_docx_renders_long_variable_as_math(tmp_path):
    """DOCX 应将长字母变量转为 OMML，而非残留字面 ``$...$``。

    Args:
        tmp_path: pytest 提供的临时目录。

    Returns:
        None.
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "Long variable $velocityvelocityvelocity$ should render as math.\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(p.text for p in document.paragraphs)
    # 长变量名应转为 OMML，不应残留字面 $...$
    assert "$velocityvelocityvelocity$" not in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml


def test_process_text_inline_logs_warning_on_latex_conversion_failure(caplog, monkeypatch):
    """``_process_text_inline`` 在 LaTeX→OMML 转换失败时应记录 warning 日志。

    回归测试：commit 867d0e6 将 ``except ValueError`` 放宽为 ``except Exception``
    但未添加任何日志记录，导致转换失败被静默吞掉。
    """
    import logging

    from openjiuwen_deepsearch.algorithm.report_export import word_utils
    from openjiuwen_deepsearch.algorithm.report_export.word_utils import (
        _process_text_inline,
    )

    def _raise_value_error(_latex: str) -> str:
        raise ValueError("mocked latex2omml failure")

    monkeypatch.setattr(word_utils, "_latex_to_omml", _raise_value_error)

    doc = Document()
    paragraph = doc.add_paragraph()

    with caplog.at_level(
        logging.WARNING,
        logger="server.deepsearch.core.manager.report_manager.word_utils",
    ):
        _process_text_inline(paragraph, "$x^2$", None)

    assert any(
        "latex" in record.message.lower() or "omml" in record.message.lower()
        for record in caplog.records
    ), f"Expected warning log for LaTeX conversion failure, got: {[r.message for r in caplog.records]}"


def test_add_hyperlink_logs_warning_on_latex_conversion_failure(caplog, monkeypatch):
    """``_add_hyperlink`` 在 LaTeX→OMML 转换失败时应记录 warning 日志。

    回归测试：commit 867d0e6 在超链接公式处理路径中也使用了无日志的
    ``except Exception``，导致转换失败被静默吞掉。
    """
    import logging

    from openjiuwen_deepsearch.algorithm.report_export import word_utils
    from openjiuwen_deepsearch.algorithm.report_export.word_utils import (
        _add_hyperlink,
    )

    def _raise_value_error(_latex: str) -> str:
        raise ValueError("mocked latex2omml failure")

    monkeypatch.setattr(word_utils, "_latex_to_omml", _raise_value_error)

    doc = Document()
    paragraph = doc.add_paragraph()

    with caplog.at_level(
        logging.WARNING,
        logger="server.deepsearch.core.manager.report_manager.word_utils",
    ):
        _add_hyperlink(paragraph, "https://example.com", "Formula $x^2$ here")

    assert any(
        "latex" in record.message.lower() or "omml" in record.message.lower()
        for record in caplog.records
    ), f"Expected warning log for LaTeX conversion failure, got: {[r.message for r in caplog.records]}"


def test_is_currency_start_recognizes_comparison_and_multiplication_operators():
    """货币前缀后跟 <、>、×、÷ 等数学运算符不应判为货币。

    回归测试：commit 867d0e6 的 _MATH_FEATURE_AFTER_CURRENCY_RE 与
    _MATH_CONTINUATION_RE 均遗漏 <、>、×、÷ 等数学特征，导致
    ``$4 < x$``、``$4 × 5$`` 被误判为货币，不进入公式扫描。
    """
    from openjiuwen_deepsearch.algorithm.report_export.conversion_utils import (
        _is_currency_start,
    )

    # $4 < x$ — < 是数学比较运算符，应判为非货币（即应进入公式扫描）
    assert _is_currency_start("$4 < x$", 0) is False
    # $4 > x$ — > 是数学比较运算符
    assert _is_currency_start("$4 > x$", 0) is False
    # $4 × 5$ — × 是数学乘法运算符
    assert _is_currency_start("$4 × 5$", 0) is False
    # $4 ÷ 2$ — ÷ 是数学除法运算符
    assert _is_currency_start("$4 ÷ 2$", 0) is False
    # $4<x$ — 紧贴无空格也应命中
    assert _is_currency_start("$4<x$", 0) is False
    # $4×5$ — 紧贴无空格也应命中
    assert _is_currency_start("$4×5$", 0) is False

    # 真正的货币仍应判为货币
    assert _is_currency_start("$4 和 $5", 0) is True
    assert _is_currency_start("$1,200.50", 0) is True


def test_protect_math_spans_keeps_comparison_math_with_currency_prefix():
    """protect_math_spans 应保护 $4 < x$、$4 × 5$ 等比较/乘法公式。

    回归测试：commit 867d0e6 后这些公式被 _is_currency_start 误判为货币，
    不进入公式扫描，DOCX 实测不含 OMML。
    """
    text = "Compare $4 < x$ and multiply $4 × 5$."

    _, formulas = protect_math_spans(text)

    assert "$4 < x$" in formulas
    assert "$4 × 5$" in formulas


def test_convert_md_to_docx_renders_comparison_and_multiplication_math(tmp_path):
    """DOCX 应将 $4 < x$、$4 × 5$ 等比较/乘法公式转为 OMML。

    回归测试：commit 867d0e6 后 _is_currency_start 遗漏 <、× 等数学特征，
    导致这些公式不进公式扫描，DOCX 不含 OMML。
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "Compare $4 < x$ and multiply $4 × 5$.\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(p.text for p in document.paragraphs)
    assert "$4 < x$" not in paragraph_text
    assert "$4 × 5$" not in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml


def test_report_html_convert_from_markdown_includes_currency_protection_script(tmp_path):
    """convert_md_to_html 输出应包含货币美元保护脚本。

    回归测试：commit 867d0e6 的 _katex_resources 缺少 escapeCurrencyDollars，
    导致 $4 和 $5 在浏览器中被 KaTeX auto-render 误配对为公式并吞掉
    两个美元符号。

    注：原 ReportHtml 类已在 report_export 重构中删除，本测试改为直接调用
    convert_md_to_html 验证 HTML 模板内的 KaTeX 货币保护逻辑。
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "变量 $G$ 为组内采样数量，价格 $4 和 $5 保持文本。\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)
    html_text = html_path.read_text(encoding="utf-8")

    # 货币保护 JS 应存在
    assert "escapeCurrency" in html_text
    # 全角美元符号占位符应存在
    assert "\uFF04" in html_text
    # $G$ 应保留为公式定界符
    assert "$G$" in html_text
    # $4 和 $5 应保留为文本（不被 KaTeX 配对）
    assert "$4 和 $5" in html_text


def test_iter_math_spans_yields_block_math_with_numeric_content():
    """_iter_math_spans 应产出 $$1$$、$$2026$$ 等纯数字块级公式。

    回归测试：commit 867d0e6 的 _iter_math_spans 对块级 $$...$$ 仍要求
    _is_likely_inline_math 返回真，而该函数拒绝纯数字内容，导致
    $$1$$ 和 $$2026$$ 在 DOCX 中保留为字面文本而非 OMML。
    """
    from openjiuwen_deepsearch.algorithm.report_export.word_utils import (
        _iter_math_spans,
    )

    # $$1$$ — 单数字块级公式
    spans = list(_iter_math_spans("text $$1$$ end"))
    assert spans == [(5, 10)]

    # $$2026$$ — 多位数字
    spans = list(_iter_math_spans("year $$2026$$ end"))
    assert spans == [(5, 13)]

    # $$E = mc^2$$ — 含等式的块级公式仍正常工作
    spans = list(_iter_math_spans("formula $$E = mc^2$$ end"))
    assert len(spans) == 1
    assert spans[0] == (8, 20)


def test_convert_md_to_docx_renders_block_math_with_numeric_content(tmp_path):
    """DOCX 应将 $$1$$、$$2026$$ 等纯数字块级公式转为 OMML。

    回归测试：commit 867d0e6 的 _iter_math_spans 拒绝纯数字块级公式，
    导致 DOCX 中保留字面 $$1$$、$$2026$$ 而非 OMML。
    """
    md_path = tmp_path / "report.md"
    docx_path = tmp_path / "report.docx"
    md_path.write_text(
        "# 测试\n\n"
        "Year $$2026$$ and count $$1$$.\n",
        encoding="utf-8",
    )

    convert_md_to_docx(md_path, docx_path)

    document = Document(docx_path)
    paragraph_text = "\n".join(p.text for p in document.paragraphs)
    assert "$$2026$$" not in paragraph_text
    assert "$$1$$" not in paragraph_text
    with zipfile.ZipFile(docx_path) as zip_file:
        document_xml = zip_file.read("word/document.xml").decode("utf-8")
    assert "<m:oMath" in document_xml


def test_report_html_currency_protection_math_ops_synced_with_backend(tmp_path):
    """HTML 货币保护脚本的 MATH_OPS 应与后端 _is_currency_start 的数学特征集同步。

    回归测试：commit 688516d 的 _is_currency_start 已将 <、>、×、÷ 当作
    数学特征（_MATH_CONTINUATION_RE / _MATH_FEATURE_AFTER_CURRENCY_RE），
    但 HTML 路径的 JS MATH_OPS 缺少这些字符，导致 $4 < x$、$4 × 5$
    在 HTML 中先被替换为货币占位符，KaTeX 不渲染，最终只显示字面公式，
    而 DOCX 则会渲染，造成两种导出不一致。
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text("价格 $4 < x$ 和 $4 × 5$ 保持公式。\n", encoding="utf-8")

    convert_md_to_html(md_path, html_path)
    html_text = html_path.read_text(encoding="utf-8")

    # MATH_OPS 应包含 <、>、×、÷，与后端 _MATH_FEATURE_AFTER_CURRENCY_RE 一致
    math_ops_match = re.search(r"var\s+MATH_OPS\s*=\s*'([^']*)'", html_text)
    assert math_ops_match is not None, "MATH_OPS 变量应存在于 HTML 脚本中"
    math_ops = math_ops_match.group(1)
    for operator in ("<", ">", "×", "÷"):
        assert operator in math_ops, (
            f"MATH_OPS 应包含 '{operator}' 以与后端 _is_currency_start 同步，"
            f"实际 MATH_OPS={math_ops!r}"
        )


def test_convert_md_to_html_keeps_nested_list_level_across_chart_block(tmp_path):
    """图片插入嵌套列表项之间时，不应打断嵌套列表层级。

    回归测试：normalize_interrupted_nested_list_blocks 应将图片和图注
    重新缩进到父列表项内部，保证后续同级列表项不被提升为顶级列表。
    原 UT 在 commit 9d72e46 删除函数时一并删除，此处恢复。
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "- **开源与闭源博弈的多维透视**：\n"
        "    - **地缘维度**：第一条\n"
        "\n"
        "![日本Top10](chart.png)\n"
        "<font size=2>**日本Top10**: 图注</font>\n"
        "    - **生态维度**：第二条\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    parent_item = next(
        item
        for item in soup.find_all("li")
        if "开源与闭源博弈的多维透视" in item.get_text()
    )
    nested_items = parent_item.find("ul", recursive=False).find_all("li", recursive=False)
    assert len(nested_items) == 2
    assert "地缘维度" in nested_items[0].get_text()
    assert "生态维度" in nested_items[1].get_text()
    assert nested_items[0].find("img", alt="日本Top10") is not None


def test_convert_md_to_html_keeps_nested_list_level_across_font_description(tmp_path):
    """font 描述块出现在列表项之间时，不应将后续嵌套列表项提升层级。

    回归测试：normalize_interrupted_nested_list_blocks 应将 font 描述块
    重新缩进到父列表项内部，保证后续嵌套列表项层级不丢失。
    原 UT 在 commit 9d72e46 删除函数时一并删除，此处恢复。
    """
    md_path = tmp_path / "report.md"
    html_path = tmp_path / "report.html"
    md_path.write_text(
        "- **多维度协同分析**：\n"
        "\n"
        "<font size=2>**边缘智能体多维度协同分析**: 描述</font>\n"
        "    - **技术维度**：内容。\n"
        "    - **经济维度**：内容。\n",
        encoding="utf-8",
    )

    convert_md_to_html(md_path, html_path)

    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    parent_item = next(
        item for item in soup.find_all("li") if "多维度协同分析" in item.get_text()
    )
    nested_items = parent_item.find("ul", recursive=False).find_all("li", recursive=False)
    assert [item.get_text(strip=True) for item in nested_items] == [
        "技术维度：内容。",
        "经济维度：内容。",
    ]
