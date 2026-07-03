# -*- coding: UTF-8 -*-
# Copyright (c) Huawei Technologies Co., Ltd. 2025-2025. All rights reserved.
"""Shared helpers for report export conversions."""

from __future__ import annotations

import html
import logging
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

TEXT_READ_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030", "gbk")
DEFAULT_DOCX_FONT = "Microsoft YaHei"
DOCX_IMAGE_UPSCALE_FACTOR = 3
DOCX_IMAGE_CONTRAST = 1.28
DOCX_IMAGE_SHARPNESS = 1.45
DOCX_IMAGE_COLOR = 1.08
SAFE_FILENAME_RE = re.compile(r"[^\w.-]+", re.UNICODE)
NUMBERED_HEADING_RE = re.compile(
    r"^(?P<indent>\s{0,3})(?P<number>\d+(?:\.\d+)*)(?:\.\s+|\s+)(?P<title>.+?)\s*$"
)
LIST_ITEM_RE = re.compile(r"^\s{0,3}(?:[-*+]\s+|\d+\.\s+)")
# Math formula spans. Block math ($$...$$) is matched first (DOTALL) so inline
# matching never splits a block delimiter pair.
BLOCK_MATH_RE = re.compile(r"\$\$.+?\$\$", re.DOTALL)
SIMPLE_VARIABLE_MATH_RE = re.compile(r"[A-Za-zΑ-Ωα-ω]")
VARIABLE_ATOM_PATTERN = r"(?:\\[A-Za-z]+|[A-Za-zΑ-Ωα-ω](?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)?)"
COMPACT_VARIABLE_MATH_RE = re.compile(r"[A-Za-zΑ-Ωα-ω]{2,}(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)*")
ALNUM_VARIABLE_MATH_RE = re.compile(r"[A-Za-zΑ-Ωα-ω]+\d+(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)*")
VARIABLE_LIST_ELLIPSIS_PATTERN = r"(?:\.{3}|…|\\ldots|\\cdots)"
VARIABLE_LIST_PATTERN = (
    rf"{VARIABLE_ATOM_PATTERN}(?:\s*,\s*{VARIABLE_ATOM_PATTERN})+"
    rf"(?:\s*,\s*{VARIABLE_LIST_ELLIPSIS_PATTERN})?"
)
VARIABLE_LIST_MATH_RE = re.compile(VARIABLE_LIST_PATTERN)
PAREN_VARIABLE_LIST_MATH_RE = re.compile(
    rf"\(\s*{VARIABLE_LIST_PATTERN}\s*\)"
)
MATH_NAME_PATTERN = (
    r"(?:\\?[A-Za-zΑ-Ωα-ω]+|[0-9]+|\\[A-Za-z]+)"
    r"(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?|\\[A-Za-z]+)*"
)
MATH_FUNCTION_CALL_PATTERN = rf"{MATH_NAME_PATTERN}\s*[\(\[][^)\]\n]+[\)\]]"
MATH_SPACING_PATTERN = r"(?:\s+|\\[,;:! ]|\\quad|\\qquad)"
MATH_DELIMITED_CONTENT_PATTERN = r"[^|\n]+"
MATH_FEATURE_RE = re.compile(r"(\\[A-Za-z]+|[_^={}]|[+\-*/!×÷∑∫√≈≠≤≥<>])")
STRONG_NUMERIC_MATH_FEATURE_RE = re.compile(r"(\\[A-Za-z]+|[_^{}]|[+\-*/!−×÷∑∫√=≈≠≤≥<>])")
MATH_FUNCTION_CALL_RE = re.compile(MATH_FUNCTION_CALL_PATTERN)
PRIME_FUNCTION_CALL_RE = re.compile(
    rf"{MATH_NAME_PATTERN}'{{1,3}}\s*[\(\[][^)\]\n]+[\)\]]"
)
FUNCTION_CALL_SEQUENCE_RE = re.compile(
    rf"{MATH_FUNCTION_CALL_PATTERN}(?:{MATH_SPACING_PATTERN}{MATH_FUNCTION_CALL_PATTERN})+"
)
ABSOLUTE_VALUE_MATH_RE = re.compile(rf"\|{MATH_DELIMITED_CONTENT_PATTERN}\|")
NORM_MATH_RE = re.compile(
    rf"(?:\|\|{MATH_DELIMITED_CONTENT_PATTERN}\|\||\\\|{MATH_DELIMITED_CONTENT_PATTERN}\\\||\\lVert\s*.+?\s*\\rVert)"
)
PERCENT_MATH_RE = re.compile(rf"(?:[+-]?\d+(?:\.\d+)?|{MATH_NAME_PATTERN})\s*\\?%")
EQUATION_REFERENCE_RE = re.compile(r"\((?:Eq\.?|Equation)\s*\d+[A-Za-z]?\)", re.IGNORECASE)
NUMERIC_TUPLE_RE = re.compile(r"\(\s*[+-]?\d+(?:\.\d+)?(?:\s*,\s*[+-]?\d+(?:\.\d+)?)+\s*\)")
PRIME_VARIABLE_RE = re.compile(
    r"(?:\\[A-Za-z]+|[A-Za-zΑ-Ωα-ω])(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)*'+"
)
PLAIN_CURRENCY_RE = re.compile(r"\d{1,3}(?:,\d{3})*(?:\.\d+)?(?:\s*(?:[-–—~至到]|to)\s*\d{1,3}(?:,\d{3})*(?:\.\d+)?)?")
# Sentinel uses NUL control chars so Python-Markdown leaves it untouched (no
# emphasis, strong, escaping, or link parsing applied to the placeholder).
MATH_PLACEHOLDER = "\x00MATH{}\x00"
MATH_PLACEHOLDER_RE = re.compile(r"\x00MATH(\d+)\x00")
INLINE_MATH_DOLLAR_RE = re.compile(r"^\$(?!\$)(.*?)(?<!\$)\$$", re.DOTALL)
HTML_TAG_RE = re.compile(r"</?[A-Za-z][A-Za-z0-9:-]*(?:\s+[^<>]*)?>")
MARKDOWN_LINK_RE = re.compile(r"!?\[[^\]\n]+\]\([^)]+\)")
CODE_PLACEHOLDER = "\x00CODE{}\x00"
CODE_PLACEHOLDER_RE = re.compile(r"\x00CODE(\d+)\x00")
FENCED_CODE_RE = re.compile(
    r"(?ms)^(?P<indent>[ \t]{0,3})(?P<fence>`{3,}|~{3,})[^\n]*\n.*?\n(?P=indent)(?P=fence)[ \t]*$"
)
INLINE_CODE_RE = re.compile(r"`+[^`\n]*`+")
LIST_ITEM_WITH_INDENT_RE = re.compile(
    r"^(?P<indent>[ \t]*)(?P<marker>(?:[-*+]\s+|\d+\.\s+).*)$"
)
INDENTED_LIST_ITEM_RE = re.compile(r"^(?P<indent>[ \t]{4,})(?P<marker>(?:[-*+]\s+|\d+\.\s+).*)$")
MARKDOWN_IMAGE_LINE_RE = re.compile(r"^[ \t]*!\[[^\]]*\]\([^)]+\)[ \t]*$")
LEGACY_FONT_CAPTION_PREFIX_RE = re.compile(
    r'^[ \t]*<font\b[^>]*\bsize\s*=\s*["\']?2["\']?[^>]*>',
    flags=re.IGNORECASE,
)
INDENTED_HTML_BLOCK_END_RE = re.compile(
    r"^[ \t]+</(?:div|section|article|main)>[ \t]*$",
    flags=re.IGNORECASE,
)
MARKDOWN_TABLE_ROW_RE = re.compile(r"^[ \t]{0,3}\|")
MARKDOWN_TABLE_DELIMITER_RE = re.compile(r":?-{1,}:?")
SENTENCE_END_RE = re.compile(r"[。！？?!…]$")
CITATION_RE = re.compile(r"\[\[(\d+)\]\]\((https?://[^\s)]+(?:\([^\s)]+\)[^\s)]*)*)\)")
CHECKED_CITATION_RE = re.compile(
    r"\[\s*checked_citation:\s*\d+\s*\](\[\[\d+\]\]\((?:[^()]|\([^()]*\))*\))"
)
LEGACY_CITATION_RE = re.compile(r"\[\s*citation:\s*\d+\s*\]")
REFERENCE_LINE_RE = re.compile(r"^(?P<indent>\s*)\[(\d+)\]\.\s+(.*)$", re.MULTILINE)
CENTER_CAPTION_RE = re.compile(r'<div\s+style="text-align:\s*center;?">', flags=re.IGNORECASE)
LEGACY_FONT_CAPTION_LINE_RE = re.compile(
    r'^(?P<indent>[ \t]*)<font\b[^>]*\bsize\s*=\s*["\']?2["\']?[^>]*>(?P<body>.*?)</font>'
    r'(?P<citations>(?:<sup class="citation">.*?</sup>)*)[ \t]*$',
    flags=re.IGNORECASE | re.MULTILINE,
)
EXTERNAL_LINK_RE = re.compile(
    r'<a\s+([^>]*?)href="(https?://[^"]+)"(?![^>]*\btarget=)([^>]*)>',
    flags=re.IGNORECASE,
)
CITATION_ANCHOR_RE = re.compile(
    r'(?<!<sup class="citation">)(<a\b[^>]*href="https?://[^"]+"[^>]*>\[(\d+)\]</a>)',
    flags=re.IGNORECASE,
)
HTML_TABLE_RE = re.compile(r"<table\b[^>]*>.*?</table>", flags=re.IGNORECASE | re.DOTALL)
TABLE_WRAP_OPEN_RE = re.compile(
    r'<div\b[^>]*\bclass=["\'][^"\']*\btable-wrap\b[^"\']*["\'][^>]*>\s*$',
    flags=re.IGNORECASE,
)
DOCX_TABLE_CAPTION_RE = re.compile(
    r"^(?:表\s*[\d一二三四五六七八九十]+(?:[-－—.][\d一二三四五六七八九十]+)*|"
    r"Table\s+[\w]+(?:[-－—.][\w]+)*)\s*[:：]",
    flags=re.IGNORECASE,
)

try:
    from PIL import Image, ImageEnhance

    PIL_AVAILABLE = True
except Exception:  # pragma: no cover - optional dependency branch
    PIL_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_STYLE_AVAILABLE = True
except Exception:  # pragma: no cover - optional dependency branch
    DOCX_STYLE_AVAILABLE = False


@dataclass(slots=True)
class MermaidRenderStats:
    """Collect Mermaid rendering statistics for logging.

    Attributes:
        total: Mermaid 代码块总数。
        success: 成功渲染的 Mermaid 数量。
        failed: 渲染失败并回退源码块的数量。
    """

    total: int = 0
    success: int = 0
    failed: int = 0


def read_text_with_fallback(path: Path) -> str:
    """Read a text file with a short fallback encoding list.

    Args:
        path: 需要读取的文本文件路径。

    Returns:
        str: 解码后的文本内容。

    Raises:
        UnicodeDecodeError: 所有候选编码都失败时抛出。
    """
    last_error: UnicodeDecodeError | None = None
    for encoding in TEXT_READ_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc

    raise UnicodeDecodeError(
        getattr(last_error, "encoding", "unknown"),
        getattr(last_error, "object", b""),
        getattr(last_error, "start", 0),
        getattr(last_error, "end", 0),
        f"Unable to decode text file: {path}",
    )


def make_safe_filename_component(value: str, *, default: str = "document") -> str:
    """Normalize a string into a filesystem-safe filename fragment.

    Args:
        value: 原始文件名片段。
        default: 归一化后为空时使用的默认值。

    Returns:
        str: 可安全用于文件名的字符串。
    """
    cleaned = SAFE_FILENAME_RE.sub("_", value).strip("._")
    return cleaned or default


def normalize_whitespace(text: str) -> str:
    """Normalize newlines and common spacing issues in Markdown text.

    Args:
        text: 原始文本内容。

    Returns:
        str: 归一化后的文本内容。
    """
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.replace("\u00a0", " ").replace("\u3000", " ")


def replace_citations(text: str) -> str:
    """Convert citation markdown into HTML superscript links.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 处理引用格式后的文本。
    """

    def _replace(match: re.Match[str]) -> str:
        idx, url = match.group(1), match.group(2).strip()
        safe_url = html.escape(url, quote=True)
        return (
            f'<sup class="citation">'
            f'<a href="{safe_url}" target="_blank" rel="noopener noreferrer">[{idx}]</a>'
            f"</sup>"
        )

    text = CITATION_RE.sub(_replace, text)
    return re.sub(r'[ \t]+(<sup class="citation">)', r"\1", text)


def strip_internal_citation_markers(text: str) -> str:
    """移除仅供内部处理使用的 citation 控制标记。

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 清理内部标记后的 Markdown 文本。
    """
    text = CHECKED_CITATION_RE.sub(r"\1", text)
    return LEGACY_CITATION_RE.sub("", text)


def normalize_reference_lines(text: str) -> str:
    """Normalize numbered reference lines into bullet items.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 归一化后的 Markdown 文本。
    """
    lines = text.splitlines()
    result: list[str] = []
    in_reference_section = False

    for line in lines:
        stripped = line.strip().lower()
        if stripped in {
            "# 参考文献",
            "## 参考文献",
            "### 参考文献",
            "# references",
            "## references",
            "### references",
        }:
            in_reference_section = True
            result.append(line)
            continue

        if in_reference_section:
            match = REFERENCE_LINE_RE.match(line)
            if match:
                indent = match.group("indent")
                idx = match.group(2)
                content = match.group(3)
                result.append(f"{indent}- [{idx}] {content}")
                continue

            if stripped == "" or re.match(r"^\s*[-*]\s+", line):
                result.append(line)
                continue

            if re.match(r"^\s*#{1,6}\s+", line):
                in_reference_section = False

        result.append(line)

    return "\n".join(result)


def fix_center_caption_blocks(text: str) -> str:
    """Rewrite centered caption HTML blocks into markdown-aware containers.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正后的 Markdown 文本。
    """
    return CENTER_CAPTION_RE.sub('<div class="figure-caption" markdown="1">', text)


def normalize_legacy_font_caption_blocks(text: str) -> str:
    """Rewrite legacy font-based captions into block-level caption containers.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正后的 Markdown 文本。
    """

    def _replace(match: re.Match[str]) -> str:
        indent = match.group("indent")
        body = match.group("body").strip()
        citations = match.group("citations").strip()
        caption_content = f"{body}{citations}"
        return (
            f'\n{indent}<div class="figure-caption" markdown="1">\n'
            f"{indent}{caption_content}\n"
            f"{indent}</div>\n"
        )

    return LEGACY_FONT_CAPTION_LINE_RE.sub(_replace, text)


def normalize_interrupted_nested_list_blocks(text: str) -> str:
    """Keep image or description blocks inside the list items they interrupt."""
    lines = text.split("\n")

    def _indent_width(indent: str) -> int:
        return len(indent.expandtabs(4))

    for index, line in enumerate(lines):
        current_item = LIST_ITEM_WITH_INDENT_RE.match(line)
        if current_item is None:
            continue

        list_indent = _indent_width(current_item.group("indent"))
        block_indexes: list[int] = []
        saw_image = False
        saw_font_description = False

        for cursor in range(index + 1, len(lines)):
            candidate = lines[cursor]
            if not candidate.strip():
                block_indexes.append(cursor)
                continue

            if MARKDOWN_IMAGE_LINE_RE.match(candidate):
                saw_image = True
                block_indexes.append(cursor)
                continue

            if saw_image and LEGACY_FONT_CAPTION_PREFIX_RE.match(candidate):
                saw_font_description = True
                block_indexes.append(cursor)
                continue

            if not saw_image and LEGACY_FONT_CAPTION_PREFIX_RE.match(candidate):
                saw_font_description = True
                block_indexes.append(cursor)
                continue

            next_item = LIST_ITEM_WITH_INDENT_RE.match(candidate)
            next_indent = _indent_width(next_item.group("indent")) if next_item else -1
            is_matching_sibling = (
                saw_image
                and next_item is not None
                and next_indent == list_indent
            )
            is_nested_after_description = (
                saw_font_description
                and next_item is not None
                and next_indent > list_indent
            )
            if is_matching_sibling or is_nested_after_description:
                content_indent_width = next_indent if is_nested_after_description else list_indent + 4
                content_indent = " " * content_indent_width
                for block_index in block_indexes:
                    if lines[block_index].strip():
                        lines[block_index] = content_indent + lines[block_index].lstrip()
            break

    return "\n".join(lines)


def normalize_list_boundaries(text: str) -> str:
    """Insert a blank line before list items when the source omits one.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正列表边界后的 Markdown 文本。
    """
    block_prefixes = ("#", ">", "|", "```", "<div", "<ul", "<ol", "<li")

    def _needs_blank_line_before_list(current_line: str, previous_line: str) -> bool:
        previous = previous_line.strip()
        if not LIST_ITEM_RE.match(current_line):
            return False
        if not previous:
            return False
        if LIST_ITEM_RE.match(previous_line):
            return False
        return not previous.startswith(block_prefixes)

    lines = text.split("\n")
    normalized: list[str] = []

    for line in lines:
        previous_line = normalized[-1] if normalized else ""

        if _needs_blank_line_before_list(line, previous_line):
            normalized.append("")

        normalized.append(line)

    return "\n".join(normalized)


def normalize_table_boundaries(text: str) -> str:
    """Insert a blank line before Markdown pipe tables when the source omits one.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正表格边界后的 Markdown 文本。
    """

    def _is_table_row(line: str) -> bool:
        return bool(MARKDOWN_TABLE_ROW_RE.match(line))

    def _is_table_delimiter(line: str) -> bool:
        if not _is_table_row(line):
            return False
        cells = [
            cell.strip().replace(" ", "")
            for cell in line.strip().strip("|").split("|")
        ]
        return len(cells) >= 2 and all(
            bool(MARKDOWN_TABLE_DELIMITER_RE.fullmatch(cell))
            for cell in cells
        )

    lines = text.split("\n")
    normalized: list[str] = []
    in_fenced_code = False

    for index, line in enumerate(lines):
        if line.lstrip().startswith("```"):
            in_fenced_code = not in_fenced_code
            normalized.append(line)
            continue

        is_table_start = (
            not in_fenced_code
            and index + 1 < len(lines)
            and _is_table_row(line)
            and _is_table_delimiter(lines[index + 1])
        )
        if is_table_start and normalized and normalized[-1].strip():
            normalized.append("")

        normalized.append(line)

    return "\n".join(normalized)


def normalize_orphan_indented_list_items(text: str) -> str:
    """Dedent report-style list items that Markdown would otherwise treat as code.

    Args:
        text: Raw Markdown text.

    Returns:
        str: Markdown text with orphan indented list items dedented.
    """
    lines = text.split("\n")
    normalized: list[str] = []
    in_fenced_code = False
    orphan_list_indent: int | None = None

    def _previous_nonempty_line() -> str:
        for previous in reversed(normalized):
            if previous.strip():
                return previous
        return ""

    def _indent_width(indent: str) -> int:
        return len(indent.expandtabs(4))

    def _line_indent_width(line: str) -> int:
        indent = re.match(r"^[ \t]*", line).group(0)
        return _indent_width(indent)

    for line in lines:
        if line.lstrip().startswith("```"):
            in_fenced_code = not in_fenced_code
            normalized.append(line)
            orphan_list_indent = None
            continue

        match = INDENTED_LIST_ITEM_RE.match(line)
        if match and not in_fenced_code:
            indent_width = _indent_width(match.group("indent"))
            if orphan_list_indent == indent_width:
                normalized.append(match.group("marker"))
                continue

            previous = _previous_nonempty_line()
            previous_indent_width = _line_indent_width(previous)
            follows_nested_content = (
                previous_indent_width > indent_width
                or (
                    previous_indent_width == indent_width
                    and INDENTED_HTML_BLOCK_END_RE.match(previous)
                )
            )
            if (
                not follows_nested_content
                and not LIST_ITEM_RE.match(previous)
                and not INDENTED_LIST_ITEM_RE.match(previous)
            ):
                orphan_list_indent = indent_width
                normalized.append(match.group("marker"))
                continue

            orphan_list_indent = None
        elif line.strip():
            orphan_list_indent = None

        normalized.append(line)

    return "\n".join(normalized)


def render_mermaid_supplement(supplement_markdown: str) -> str:
    """Render timeline supplement markdown into an HTML helper block.

    Args:
        supplement_markdown: Mermaid 预处理阶段生成的补充说明 Markdown。

    Returns:
        str: 内嵌 HTML 容器字符串。
    """
    supplement_markdown = supplement_markdown.strip()
    if not supplement_markdown:
        return ""
    return (
        '\n<div class="timeline-notes" markdown="1">\n'
        f"{supplement_markdown}\n"
        "</div>\n"
    )


def preprocess_markdown_text(text: str) -> str:
    """Apply Markdown preprocessing before HTML or DOCX conversion.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 预处理后的 Markdown 文本。
    """
    text = normalize_whitespace(text)
    text = normalize_interrupted_nested_list_blocks(text)
    text = strip_internal_citation_markers(text)
    text = replace_citations(text)
    text = normalize_reference_lines(text)
    text = fix_center_caption_blocks(text)
    text = normalize_legacy_font_caption_blocks(text)
    text = normalize_orphan_indented_list_items(text)
    text = normalize_list_boundaries(text)
    return normalize_table_boundaries(text)


def wrap_html_tables(html_text: str) -> str:
    """Wrap HTML tables with a scrollable centering container."""
    if "<table" not in html_text.lower():
        return html_text

    def _replace(match: re.Match[str]) -> str:
        prefix = html_text[max(0, match.start() - 512): match.start()]
        if TABLE_WRAP_OPEN_RE.search(prefix):
            return match.group(0)
        return f'<div class="table-wrap">{match.group(0)}</div>'

    return HTML_TABLE_RE.sub(_replace, html_text)


def postprocess_html(html_text: str) -> str:
    """Postprocess generated HTML for external links, citations and table wrappers.

    Args:
        html_text: Markdown 转换后的 HTML 文本。

    Returns:
        str: 处理后的 HTML 文本。
    """

    def _replace(match: re.Match[str]) -> str:
        before = match.group(1).rstrip()
        href = re.sub(r"\s+", "", match.group(2))
        after = match.group(3).rstrip()
        attrs = " ".join(part for part in [before, f'href="{href}"', after] if part)
        return f'<a {attrs} target="_blank" rel="noopener noreferrer">'

    def _wrap_citation(match: re.Match[str]) -> str:
        return f'<sup class="citation">{match.group(1)}</sup>'

    html_text = EXTERNAL_LINK_RE.sub(_replace, html_text)
    html_text = CITATION_ANCHOR_RE.sub(_wrap_citation, html_text)
    html_text = re.sub(r'[ \t]+(<sup class="citation">)', r"\1", html_text)
    html_text = re.sub(r'(</sup>)[ \t]+(<sup class="citation">)', r"\1\2", html_text)
    return wrap_html_tables(html_text)


def protect_math_spans(text: str) -> tuple[str, list[str]]:
    """Replace LaTeX math spans with placeholders before Markdown conversion.

    Python-Markdown otherwise treats characters inside ``$...$`` / ``$$...$$``
    (notably ``_`` and ``*``) as emphasis markup and injects ``<em>``/``<strong>``
    tags into the formula, leaving MathJax with invalid LaTeX. Each span is
    swapped for a NUL-wrapped sentinel that Markdown leaves untouched, then
    restored verbatim by :func:`restore_math_spans` after conversion.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        tuple[str, list[str]]: (替换占位符后的文本, 按出现顺序保存的公式原文列表)。
    """
    formulas: list[str] = []
    code_spans: list[str] = []

    def _store_code(match: re.Match[str]) -> str:
        code_spans.append(match.group(0))
        return CODE_PLACEHOLDER.format(len(code_spans) - 1)

    def _store(match: re.Match[str]) -> str:
        formulas.append(match.group(0))
        return MATH_PLACEHOLDER.format(len(formulas) - 1)

    text = FENCED_CODE_RE.sub(_store_code, text)
    text = INLINE_CODE_RE.sub(_store_code, text)
    # Block math first (DOTALL) so inline matching never splits a ``$$`` pair.
    text = BLOCK_MATH_RE.sub(_store, text)
    text = _protect_inline_math_spans(text, formulas)

    def _restore_code(match: re.Match[str]) -> str:
        index = int(match.group(1))
        if 0 <= index < len(code_spans):
            return code_spans[index]
        return match.group(0)

    text = CODE_PLACEHOLDER_RE.sub(_restore_code, text)
    return text, formulas


def _protect_inline_math_spans(text: str, formulas: list[str]) -> str:
    """Protect inline LaTeX spans while leaving currency-like dollars intact."""
    parts: list[str] = []
    cursor = 0
    index = 0

    while index < len(text):
        start = text.find("$", index)
        if start == -1:
            break
        if _should_skip_inline_math_start(text, start):
            index = start + 1
            continue

        end = _find_inline_math_end(text, start + 1)
        if end is None:
            index = start + 1
            continue

        formula = text[start:end + 1]
        if _is_likely_inline_math(formula[1:-1].strip()):
            parts.append(text[cursor:start])
            formulas.append(formula)
            parts.append(MATH_PLACEHOLDER.format(len(formulas) - 1))
            cursor = end + 1
            index = end + 1
        else:
            index = end + 1

    if not parts:
        return text

    parts.append(text[cursor:])
    return "".join(parts)


def _find_inline_math_end(text: str, start_index: int) -> int | None:
    index = start_index
    while index < len(text):
        end = text.find("$", index)
        if end == -1:
            return None
        if "\n" in text[start_index:end]:
            return None
        if _is_valid_inline_math_end(text, start_index, end):
            return end
        index = end + 1
    return None


def _should_skip_inline_math_start(text: str, index: int) -> bool:
    """Return whether a dollar sign cannot start an inline math span."""
    if _is_escaped(text, index):
        return True
    if _is_double_dollar(text, index):
        return True
    if _is_likely_currency_start(text, index):
        return True
    return not _is_valid_inline_math_start(text, index)


def _is_valid_inline_math_end(text: str, start_index: int, end_index: int) -> bool:
    """Return whether a dollar sign can close an inline math span."""
    if _is_escaped(text, end_index):
        return False
    if _is_double_dollar(text, end_index):
        return False
    if end_index <= start_index:
        return False
    if text[end_index - 1].isspace():
        return False
    return not (end_index + 1 < len(text) and text[end_index + 1].isdigit())


def _is_likely_inline_math(content: str) -> bool:
    if not content:
        return False
    if HTML_TAG_RE.search(content) or MARKDOWN_LINK_RE.search(content):
        return False
    if PLAIN_CURRENCY_RE.fullmatch(content):
        return False
    return bool(
        SIMPLE_VARIABLE_MATH_RE.fullmatch(content)
        or COMPACT_VARIABLE_MATH_RE.fullmatch(content)
        or ALNUM_VARIABLE_MATH_RE.fullmatch(content)
        or VARIABLE_LIST_MATH_RE.fullmatch(content)
        or PAREN_VARIABLE_LIST_MATH_RE.fullmatch(content)
        or MATH_FEATURE_RE.search(content)
        or MATH_FUNCTION_CALL_RE.fullmatch(content)
        or PRIME_FUNCTION_CALL_RE.fullmatch(content)
        or FUNCTION_CALL_SEQUENCE_RE.fullmatch(content)
        or ABSOLUTE_VALUE_MATH_RE.fullmatch(content)
        or NORM_MATH_RE.fullmatch(content)
        or PERCENT_MATH_RE.fullmatch(content)
        or _is_balanced_math_function_call(content)
        or EQUATION_REFERENCE_RE.fullmatch(content)
        or NUMERIC_TUPLE_RE.fullmatch(content)
        or PRIME_VARIABLE_RE.fullmatch(content)
    )


def _is_balanced_math_function_call(content: str) -> bool:
    """Return whether content is a math-like function call with balanced args."""
    name_match = re.match(
        r"(?:\\?[A-Za-zΑ-Ωα-ω]+|[0-9]+|\\[A-Za-z]+)"
        r"\s*(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?|\\[A-Za-z]+)*\s*",
        content,
    )
    if name_match is None:
        return False

    open_index = name_match.end()
    if open_index >= len(content) or content[open_index] not in "([":
        return False

    open_char = content[open_index]
    close_char = ")" if open_char == "(" else "]"
    pairs = {"(": ")", "[": "]"}
    stack = [close_char]
    index = open_index + 1
    saw_argument = False

    while index < len(content):
        char = content[index]
        if char == "\\":
            index += 2
            saw_argument = True
            continue
        if char in pairs:
            stack.append(pairs[char])
            saw_argument = True
        elif stack and char == stack[-1]:
            stack.pop()
            if not stack:
                return saw_argument and index == len(content) - 1
        elif char in ")]":
            return False
        elif not char.isspace():
            saw_argument = True
        index += 1

    return False


def _is_valid_inline_math_start(text: str, index: int) -> bool:
    next_index = index + 1
    if next_index >= len(text):
        return False
    next_char = text[next_index]
    if next_char.isspace() or next_char in "，。；：、,.!?;:)]}）】」』":
        return False
    return True


def _is_likely_currency_start(text: str, index: int) -> bool:
    """Return whether a dollar sign begins currency-like text, not math."""
    match = re.match(r"\$(?:\d{1,3}(?:,\d{3})*|\d+)(?:\.\d+)?", text[index:])
    if not match:
        return False
    next_index = index + len(match.group(0))
    if next_index >= len(text):
        return True
    if text.startswith(r"\%", next_index) or text[next_index] == "%":
        return False
    end = _find_inline_math_end(text, index + 1)
    if end is not None and STRONG_NUMERIC_MATH_FEATURE_RE.search(text[next_index:end]):
        return False
    return text[next_index] not in "([{_=^\\"


def _is_double_dollar(text: str, index: int) -> bool:
    return (
        (index > 0 and text[index - 1] == "$")
        or (index + 1 < len(text) and text[index + 1] == "$")
    )


def _is_escaped(text: str, index: int) -> bool:
    slash_count = 0
    cursor = index - 1
    while cursor >= 0 and text[cursor] == "\\":
        slash_count += 1
        cursor -= 1
    return slash_count % 2 == 1


def restore_math_spans(html_text: str, formulas: list[str]) -> str:
    """Restore math placeholders produced by :func:`protect_math_spans`.

    Args:
        html_text: 经 Markdown 转换、仍含公式占位符的 HTML 文本。
        formulas: :func:`protect_math_spans` 返回的公式原文列表。

    Returns:
        str: 占位符还原为公式原文后的 HTML 文本。
    """

    def _restore(match: re.Match[str]) -> str:
        index = int(match.group(1))
        if 0 <= index < len(formulas):
            return normalize_math_span_for_rendering(formulas[index])
        return match.group(0)

    return MATH_PLACEHOLDER_RE.sub(_restore, html_text)


def normalize_math_span_for_rendering(formula: str) -> str:
    """Normalize protected inline dollar math to explicit LaTeX delimiters."""
    match = INLINE_MATH_DOLLAR_RE.match(formula)
    if match:
        return rf"\({html.escape(match.group(1), quote=False)}\)"
    block_match = BLOCK_MATH_RE.fullmatch(formula)
    if block_match:
        return f"$${html.escape(formula[2:-2], quote=False)}$$"
    return html.escape(formula, quote=False)


def _neighbor_numbered_line(lines: list[str], index: int, *, reverse: bool) -> str | None:
    """Find the nearest non-empty neighboring line.

    Args:
        lines: 文本行列表。
        index: 当前行索引。
        reverse: 是否向前查找。

    Returns:
        str | None: 相邻非空行文本。
    """
    step = -1 if reverse else 1
    cursor = index + step
    while 0 <= cursor < len(lines):
        stripped = lines[cursor].strip()
        if stripped:
            return stripped
        cursor += step
    return None


def _should_promote_numbered_heading(lines: list[str], index: int, match: re.Match[str]) -> bool:
    """Decide whether a numbered line should become a Markdown heading.

    Args:
        lines: 文本行列表。
        index: 当前行索引。
        match: 编号标题匹配结果。

    Returns:
        bool: 是否应提升为标题。
    """
    title = match.group("title").strip()
    if len(title) > 80 or SENTENCE_END_RE.search(title):
        return False
    if index > 0 and lines[index - 1].strip():
        return False
    if index + 1 < len(lines) and lines[index + 1].strip():
        return False

    prev_nonempty = _neighbor_numbered_line(lines, index, reverse=True)
    next_nonempty = _neighbor_numbered_line(lines, index, reverse=False)
    if prev_nonempty and NUMBERED_HEADING_RE.match(prev_nonempty):
        return False
    if next_nonempty and NUMBERED_HEADING_RE.match(next_nonempty):
        return False
    return True


def normalize_headings(content: str) -> str:
    """Normalize numbered headings into Markdown heading syntax.

    Args:
        content: 原始 Markdown 文本。

    Returns:
        str: 标题归一化后的 Markdown 文本。
    """
    content = normalize_whitespace(content)
    lines = content.split("\n")
    out: list[str] = []
    in_code_block = False

    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            out.append(line)
            continue
        if in_code_block:
            out.append(line)
            continue
        if not stripped:
            out.append("")
            continue

        match_hash = re.match(r"^(#{1,6})\s*(.+?)\s*$", stripped)
        if match_hash:
            hashes = match_hash.group(1)
            title = match_hash.group(2).strip()
            if out and out[-1] != "":
                out.append("")
            out.append(f"{hashes} {title}")
            out.append("")
            continue

        match_numbered = NUMBERED_HEADING_RE.match(line)
        if match_numbered and _should_promote_numbered_heading(lines, index, match_numbered):
            numbering = match_numbered.group("number")
            title = match_numbered.group("title").strip()
            level = min(numbering.count(".") + 1, 6)
            heading = "#" * level + " " + f"{numbering} {title}"
            if out and out[-1] != "":
                out.append("")
            out.append(heading)
            out.append("")
            continue

        out.append(line)

    result = "\n".join(out)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip() + "\n"


def _set_rfonts(rpr, font_name: str) -> None:
    """Set all Word run font slots to a target font.

    Args:
        rpr: Word run properties element。
        font_name: 目标字体名。

    Returns:
        None.
    """
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def _set_run_font(run, font_name: str) -> None:
    """Apply a font to one docx run.

    Args:
        run: python-docx run 对象。
        font_name: 目标字体。

    Returns:
        None.
    """
    run.font.name = font_name
    _set_rfonts(run.element.get_or_add_rPr(), font_name)


def _set_style_font(style, font_name: str) -> None:
    """Apply a font to one docx style.

    Args:
        style: python-docx style 对象。
        font_name: 目标字体。

    Returns:
        None.
    """
    style.font.name = font_name
    _set_rfonts(style.element.get_or_add_rPr(), font_name)


def _apply_font_to_table(table, font_name: str) -> None:
    """Apply a font recursively to a docx table.

    Args:
        table: python-docx table 对象。
        font_name: 目标字体。

    Returns:
        None.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, font_name)
            for nested_table in cell.tables:
                _apply_font_to_table(nested_table, font_name)


def _center_docx_table(table) -> None:
    """Center one docx table and its nested tables.

    Args:
        table: python-docx table 对象。

    Returns:
        None.
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for nested_table in cell.tables:
                _center_docx_table(nested_table)


def _center_docx_table_captions(paragraphs) -> None:
    """Center paragraphs that look like table captions.

    Args:
        paragraphs: python-docx paragraph iterable.

    Returns:
        None.
    """
    for paragraph in paragraphs:
        if DOCX_TABLE_CAPTION_RE.match(paragraph.text.strip()):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def normalize_docx_tables(docx_path: Path) -> None:
    """Center tables and table-caption paragraphs in a generated DOCX file.

    Args:
        docx_path: DOCX 文件路径。

    Returns:
        None.
    """
    if not DOCX_STYLE_AVAILABLE:
        logger.warning("python-docx is unavailable, skipping table normalization for %s.", docx_path)
        return

    document = Document(docx_path)
    _center_docx_table_captions(document.paragraphs)
    for table in document.tables:
        _center_docx_table(table)

    for section in document.sections:
        _center_docx_table_captions(section.header.paragraphs)
        _center_docx_table_captions(section.footer.paragraphs)
        for table in section.header.tables:
            _center_docx_table(table)
        for table in section.footer.tables:
            _center_docx_table(table)

    document.save(docx_path)


def normalize_docx_fonts(docx_path: Path, *, font_name: str = DEFAULT_DOCX_FONT) -> None:
    """Normalize the font family across a generated DOCX file.

    Args:
        docx_path: DOCX 文件路径。
        font_name: 目标字体名。

    Returns:
        None.
    """
    if not DOCX_STYLE_AVAILABLE:
        logger.warning("python-docx is unavailable, skipping font normalization for %s.", docx_path)
        return

    document = Document(docx_path)
    for style_name in (
        "Normal",
        "Title",
        "Subtitle",
        "Body Text",
        "Hyperlink",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "Heading 7",
        "Heading 8",
        "Heading 9",
    ):
        try:
            _set_style_font(document.styles[style_name], font_name)
        except KeyError:
            continue

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font_name)

    for table in document.tables:
        _apply_font_to_table(table, font_name)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for table in section.header.tables:
            _apply_font_to_table(table, font_name)
        for table in section.footer.tables:
            _apply_font_to_table(table, font_name)

    document.save(docx_path)


def enhance_image(image_path: str) -> None:
    """Upscale and sharpen Mermaid PNG images for DOCX output.

    Args:
        image_path: PNG 图片路径。

    Returns:
        None.
    """
    if not PIL_AVAILABLE:
        return

    try:
        with Image.open(image_path) as original:
            if original.mode in ("RGBA", "LA"):
                background = Image.new("RGBA", original.size, (255, 255, 255, 255))
                background.alpha_composite(original.convert("RGBA"))
                image = background.convert("RGB")
            else:
                image = original.convert("RGB")

        image = image.resize(
            (image.width * DOCX_IMAGE_UPSCALE_FACTOR, image.height * DOCX_IMAGE_UPSCALE_FACTOR),
            Image.Resampling.LANCZOS,
        )
        image = ImageEnhance.Contrast(image).enhance(DOCX_IMAGE_CONTRAST)
        image = ImageEnhance.Sharpness(image).enhance(DOCX_IMAGE_SHARPNESS)
        image = ImageEnhance.Color(image).enhance(DOCX_IMAGE_COLOR)
        image.save(image_path, format="PNG", optimize=True)
    except Exception as exc:  # pragma: no cover - best effort enhancement
        logger.warning("Failed to enhance image %s: %s", image_path, exc)
