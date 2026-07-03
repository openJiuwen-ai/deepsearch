# -*- coding: UTF-8 -*-
# Copyright (c) Huawei Technologies Co., Ltd. 2025-2025. All rights reserved.
import base64
import io
import logging
import re
from dataclasses import dataclass, replace
from pathlib import Path
from urllib.parse import unquote, urlparse

from bs4 import BeautifulSoup, NavigableString
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph
from latex2mathml.converter import convert as latex2mathml_convert
from mathml2omml import convert

logger = logging.getLogger(__name__)

# NOTE:
# python-docx does not expose public APIs for a subset of low-level XML operations.
# The internal members accessed below are intentionally constrained to formatting helpers.

HYPERLINK_URI = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)  # URI for word hyperlink
OMML_URI = "http://schemas.openxmlformats.org/officeDocument/2006/math"  # URI for word omml
MAX_HTML_BLOCK_DEPTH = 100
HEADING_TAGS = frozenset(f"h{i}" for i in range(1, 9))
REMOTE_IMAGE_SCHEMES = frozenset({"http", "https"})
LATEX_TOKEN_RE = re.compile(r"(\$\$.*?\$\$|\\\(.*?\\\))", re.DOTALL)
LATEX_GROUPED_COMMANDS_WITH_POWER = frozenset({"binom", "frac"})
LATEX_NORMALIZATION_MAX_PASSES = 8
LATEX_ALIGNMENT_ENV_RE = re.compile(
    r"\\begin\{(?P<env>align\*?|aligned|split|gathered)\}"
    r"(?P<body>.*?)"
    r"\\end\{(?P=env)\}",
    re.DOTALL,
)
HTML_FORMATTING_WHITESPACE_RE = re.compile(r"[ \t]*\n[ \t]*")
DOCX_LIST_LEVELS = 9
DOCX_BULLET_SYMBOLS = ("•", "◦", "▪")


@dataclass(frozen=True)
class HtmlToDocContext:
    """Shared state for converting HTML nodes into DOCX content."""

    style_dict: dict
    base_path: Path | None = None
    max_image_width: int | None = None
    max_depth: int = MAX_HTML_BLOCK_DEPTH
    style_r_fonts: object | None = None
    current_run: object | None = None
    superscript: bool = False


@dataclass(frozen=True)
class HtmlBlockState:
    """Recursive state for converting HTML block elements."""

    depth: int = 0
    list_depth: int = 0
    list_num_id: int | None = None
    list_tag: str | None = None


def _docx_run_element(run):
    return run._element  # pylint: disable=protected-access


def _docx_run_r(run):
    return run._r  # pylint: disable=protected-access


def _docx_paragraph_p(paragraph):
    return paragraph._p  # pylint: disable=protected-access


def _docx_table_element(table):
    return table._element  # pylint: disable=protected-access


def _docx_style_element(style):
    return style._element  # pylint: disable=protected-access


def _append_word_list_level(abstract_num, level: int, *, ordered: bool) -> None:
    """Append one list level to a Word abstract numbering definition."""
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(level))

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(
        qn("w:val"),
        f"%{level + 1}." if ordered else DOCX_BULLET_SYMBOLS[level % len(DOCX_BULLET_SYMBOLS)],
    )
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(360 * (level + 1)))
    indent.set(qn("w:hanging"), "180")
    p_pr.append(indent)
    lvl.append(p_pr)

    abstract_num.append(lvl)


def _create_word_list_numbering(doc, *, ordered: bool) -> int:
    """Create a native multilevel Word list and return its numbering ID."""
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_num_id = max(abstract_ids, default=-1) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level_type)
    for level in range(DOCX_LIST_LEVELS):
        _append_word_list_level(abstract_num, level, ordered=ordered)
    numbering.insert(len(abstract_ids), abstract_num)

    num = numbering.add_num(abstract_num_id)
    return int(num.get(qn("w:numId")))


def _apply_word_list_numbering(paragraph, num_id: int, level: int) -> None:
    """Apply a native Word list numbering ID and level to one paragraph."""
    num_pr = _docx_paragraph_p(paragraph).get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, DOCX_LIST_LEVELS - 1)))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)


def _get_style_def_by_tag(tag: str) -> str:
    """
    将 HTML 标签映射为语义名称。

    参数:
        tag (str): HTML 标签名称，例如 'h1', 'p', 'table', 'div' 等。

    返回:
        str: 对应的语义名称，例如 'heading1', 'paragraph', 'table', 或 'default'。
    """
    tag = tag.lower().strip("<>/")  # 清理标签格式

    # 处理 heading 标签
    if tag.startswith("h") and tag[1:].isdigit():
        level = int(tag[1:])
        if 1 <= level <= 9:
            return f"heading{level}"

    # 特定标签映射
    tag_map = {
        "title": "title",
        "p": "paragraph",
        "table": "table"
    }

    return tag_map.get(tag, "default")


def _get_style_by_tag(tag_name, style_dict, doc, default="Normal") -> ParagraphStyle:
    """
    从指定文件中读取样式配置，并返回指定键的样式名称。

    参数：
    - tag_name: html中的tag名
    - style_dict: 样式dict
    - doc: 加载的带样式的Document
    - default: 如果找不到键时返回的默认值

    返回：
    - 样式名称字符串
    """
    style_def = _get_style_def_by_tag(tag_name)  # style_def是我们自己定义的名字，这里根据html的tag名拿到style_def
    style_name = style_dict.get(style_def, default)  # 再根据style_def从模板json中读取到对应的docx style_name
    return doc.styles[style_name]


def _apply_style_font_on_para_run(p: Paragraph, style_r_fonts) -> None:
    if style_r_fonts is None:
        return

    # make sure there is rFonts
    for run in p.runs:
        e = _docx_run_element(run)
        if e.rPr is None:
            e.insert(0, OxmlElement('w:rPr'))
        if e.rPr.rFonts is None:
            r_fonts = OxmlElement('w:rFonts')
            e.rPr.append(r_fonts)
        else:
            r_fonts = e.rPr.rFonts

        # set run font
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            val = style_r_fonts.get(qn(attr))
            if val:
                r_fonts.set(qn(attr), val)


def _apply_inline_style(run, tag_name):
    r_pr = _docx_run_r(run).get_or_add_rPr()

    if tag_name in ("strong", "b"):
        b = OxmlElement("w:b")
        r_pr.append(b)

    if tag_name in ("em", "i"):
        i = OxmlElement("w:i")
        r_pr.append(i)

    if tag_name == "u":
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

    if tag_name == "sup":
        vert_align = OxmlElement("w:vertAlign")
        vert_align.set(qn("w:val"), "superscript")
        r_pr.append(vert_align)


def _apply_style_font_on_run(run, style_r_fonts) -> None:
    if style_r_fonts is None:
        return

    e = _docx_run_element(run)
    if e.rPr is None:
        e.insert(0, OxmlElement('w:rPr'))
    if e.rPr.rFonts is None:
        r_fonts = OxmlElement('w:rFonts')
        e.rPr.append(r_fonts)
    else:
        r_fonts = e.rPr.rFonts

    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        val = style_r_fonts.get(qn(attr))
        if val:
            r_fonts.set(qn(attr), val)


def _apply_r_fonts_to_r_pr(r_pr, style_r_fonts) -> None:
    if style_r_fonts is None:
        return
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        val = style_r_fonts.get(qn(attr))
        if val:
            r_fonts.set(qn(attr), val)
    r_pr.append(r_fonts)


def _apply_superscript_to_r_pr(r_pr) -> None:
    vert_align = OxmlElement("w:vertAlign")
    vert_align.set(qn("w:val"), "superscript")
    r_pr.append(vert_align)


def _add_text_run(p, text: str, style_r_fonts, current_run=None, superscript: bool = False):
    if current_run is None:
        run = p.add_run(text)
    else:
        run = current_run
        run.add_text(text)

    _apply_style_font_on_run(run, style_r_fonts)
    if superscript:
        _apply_inline_style(run, "sup")
    return run


def _fit_inline_shape_to_width(inline_shape, max_width) -> None:
    if max_width is None or inline_shape.width <= max_width:
        return
    if inline_shape.width == 0:
        inline_shape.width = max_width
        return
    scale = max_width / inline_shape.width
    inline_shape.width = max_width
    inline_shape.height = int(inline_shape.height * scale)


def _add_hyperlink(paragraph, url, text, *, style_r_fonts=None, superscript: bool = False):
    # 创建关系 id
    part = paragraph.part
    r_id = part.relate_to(
        url,
        HYPERLINK_URI,
        is_external=True
    )

    # 创建 <w:hyperlink>
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # 创建 <w:r>
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    _apply_r_fonts_to_r_pr(r_pr, style_r_fonts)

    # 超链接样式（蓝色 + 下划线）
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)
    if superscript:
        _apply_superscript_to_r_pr(r_pr)

    r.append(r_pr)

    # 文本节点
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    hyperlink.append(r)
    _docx_paragraph_p(paragraph).append(hyperlink)


def _is_relative_to(path: Path, base_path: Path) -> bool:
    try:
        path.relative_to(base_path)
        return True
    except ValueError:
        return False


def _resolve_local_image(src: str, base_path: Path | None) -> Path | None:
    if base_path is None or not src:
        return None

    parsed = urlparse(src)
    if parsed.scheme in REMOTE_IMAGE_SCHEMES or parsed.scheme == "data":
        return None

    raw_path = unquote(parsed.path or src)
    image_path = Path(raw_path)
    base_resolved = base_path.resolve()
    if image_path.is_absolute():
        candidate = image_path.resolve()
    else:
        candidate = (base_resolved / image_path).resolve()

    if not _is_relative_to(candidate, base_resolved):
        return None
    return candidate if candidate.exists() else None


def _process_text_inline(p, text: str, style_r_fonts, current_run=None, superscript: bool = False) -> None:
    if "$" not in text and "\\(" not in text:
        _add_text_run(p, text, style_r_fonts, current_run=current_run, superscript=superscript)
        return

    cursor = 0
    reusable_run = current_run
    for match in LATEX_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            _add_text_run(
                p,
                text[cursor:match.start()],
                style_r_fonts,
                current_run=reusable_run,
                superscript=superscript,
            )
            reusable_run = None

        token = match.group(0)
        latex = (
            token[2:-2].strip()
            if token.startswith("\\(") and token.endswith("\\)")
            else token.strip("$").strip()
        )
        if latex:
            try:
                omml = _latex_to_omml(latex)
                _insert_omml(p, omml)
            except ValueError:
                # Fallback: render raw LaTeX text when conversion fails
                # (e.g., unbalanced \left...\right, unsupported commands)
                logger.warning(
                    "LaTeX-to-OMML conversion failed, falling back to raw text. "
                    "latex=%s",
                    latex[:200],
                )
                _add_text_run(
                    p,
                    match.group(0),
                    style_r_fonts,
                    current_run=reusable_run,
                    superscript=superscript,
                )
        cursor = match.end()

    if cursor < len(text):
        _add_text_run(
            p,
            text[cursor:],
            style_r_fonts,
            current_run=reusable_run,
            superscript=superscript,
        )


def _process_inline(
    p,
    node,
    context: HtmlToDocContext,
):
    """递归处理段落内的所有 inline 节点"""

    # 纯文本
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        if "\n" in text:
            if not text.strip():
                return
            text = HTML_FORMATTING_WHITESPACE_RE.sub(" ", text)

        _process_text_inline(
            p,
            text,
            context.style_r_fonts,
            current_run=context.current_run,
            superscript=context.superscript,
        )
        return

    # 图片（通常自己一个 run，和 current_run 无强关联）
    if node.name == "img":
        src = node.get("src")
        if src and src.startswith("data:image"):
            _, b64data = src.split(",", 1)
            img_bytes = base64.b64decode(b64data)
            run = p.add_run()
            inline_shape = run.add_picture(io.BytesIO(img_bytes))
            _fit_inline_shape_to_width(inline_shape, context.max_image_width)
            return

        image_path = _resolve_local_image(src or "", context.base_path)
        if image_path is not None:
            run = p.add_run()
            inline_shape = run.add_picture(str(image_path))
            _fit_inline_shape_to_width(inline_shape, context.max_image_width)
        return

    # 超链接：让 add_hyperlink 自己处理 run/样式
    if node.name == "a":
        href = node.get("href")
        text = node.get_text(strip=True)
        if href and text:
            _add_hyperlink(
                p,
                href,
                text,
                style_r_fonts=context.style_r_fonts,
                superscript=context.superscript,
            )
        return

    if node.name == "sup":
        for child in node.contents:
            _process_inline(
                p,
                child,
                replace(context, superscript=True),
            )
        return

    # inline 标签（strong, b, em, i, u, etc.）
    if node.name in ("strong", "b", "em", "i", "u"):
        # 如果已有 run，就在这个 run 上叠加样式；否则新建一个 run
        run = context.current_run or p.add_run()
        _apply_style_font_on_run(run, context.style_r_fonts)
        _apply_inline_style(run, node.name)

        for child in node.contents:
            _process_inline(
                p,
                child,
                replace(context, current_run=run),
            )
        return

    # 其他标签 → 递归处理，保持 current_run 传递
    for child in node.contents:
        _process_inline(
            p,
            child,
            context,
        )


def _add_para_and_apply_style(doc, element, context: HtmlToDocContext):
    style = _get_style_by_tag(element.name, context.style_dict, doc)
    p = doc.add_paragraph(style=style)

    style_r_pr = style.element.get_or_add_rPr()
    style_r_fonts = style_r_pr.find(qn('w:rFonts'))

    for child in element.contents:
        _process_inline(p, child, replace(context, style_r_fonts=style_r_fonts))


def _insert_omml(paragraph, omml_xml: str):
    """向段落中插入 OMML 公式"""
    wrapped_xml = f''' <root xmlns:m="{OMML_URI}"> {omml_xml} </root> '''
    try:
        # 1. 解析 OMML 字符串为 python-docx 可识别的 XML
        root = parse_xml(wrapped_xml)
        omath = root[0]  # 取出真正的 <m:oMath> 节点

        # 2. 插入到 run 中
        run = paragraph.add_run()
        _docx_run_r(run).append(omath)

    except Exception as e:
        raise ValueError("insert omml to doc failed") from e


def _latex_to_omml(latex: str) -> str:
    """
    将 LaTeX 数学公式转换为 Word 可识别的 OMML XML 字符串。
    依赖：
        pip install latex2mathml
        pip install lxml
    参数：
        latex: 纯 LaTeX 数学表达式（不含 $）
    返回：
        OMML XML 字符串，可直接插入 python-docx
    """

    try:
        latex = _normalize_latex_for_omml(latex)
        # 1. LaTeX → MathML
        mathml = latex2mathml_convert(latex)

        # 2. MathML → OMML（使用 mathml2omml-as）
        omml = convert(mathml)

        return omml
    except Exception as e:
        raise ValueError("transfer latex to omml failed") from e


def _normalize_latex_for_omml(latex: str) -> str:
    """Normalize valid LaTeX forms that mathml2omml cannot parse directly."""
    previous = _strip_latex_alignment_markers(latex)
    for _ in range(LATEX_NORMALIZATION_MAX_PASSES):
        current = _wrap_grouped_command_powers(previous)
        if current == previous:
            return current
        previous = current
    return previous


def _strip_latex_alignment_markers(latex: str) -> str:
    """Remove unescaped alignment markers from LaTeX alignment environments."""

    def _strip_environment(match: re.Match[str]) -> str:
        env = match.group("env")
        body = _strip_unescaped_latex_char(match.group("body"), "&")
        return rf"\begin{{{env}}}{body}\end{{{env}}}"

    return LATEX_ALIGNMENT_ENV_RE.sub(_strip_environment, latex)


def _strip_unescaped_latex_char(text: str, target: str) -> str:
    parts: list[str] = []
    index = 0
    while index < len(text):
        char = text[index]
        if char == "\\" and index + 1 < len(text):
            parts.append(text[index:index + 2])
            index += 2
            continue
        if char != target:
            parts.append(char)
        index += 1
    return "".join(parts)


def _wrap_grouped_command_powers(latex: str) -> str:
    parts: list[str] = []
    cursor = 0
    index = 0

    while index < len(latex):
        match = re.search(r"\\([A-Za-z]+)", latex[index:])
        if match is None:
            break

        command_start = index + match.start()
        command_end = index + match.end()
        command_name = match.group(1)
        if command_name not in LATEX_GROUPED_COMMANDS_WITH_POWER:
            index = command_end
            continue

        first_group_end = _find_latex_group_end(latex, command_end)
        if first_group_end is None:
            index = command_end
            continue
        second_group_end = _find_latex_group_end(latex, first_group_end + 1)
        if second_group_end is None:
            index = first_group_end + 1
            continue

        power_end = _find_latex_power_end(latex, second_group_end + 1)
        if power_end is None:
            index = command_end
            continue

        parts.append(latex[cursor:command_start])
        parts.append("{")
        parts.append(latex[command_start:second_group_end + 1])
        parts.append("}")
        parts.append(latex[second_group_end + 1:power_end])
        cursor = power_end
        index = power_end

    if not parts:
        return latex

    parts.append(latex[cursor:])
    return "".join(parts)


def _find_latex_group_end(text: str, open_index: int) -> int | None:
    if open_index >= len(text) or text[open_index] != "{":
        return None

    depth = 0
    index = open_index
    while index < len(text):
        char = text[index]
        if char == "\\":
            index += 2
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return index
        index += 1
    return None


def _find_latex_power_end(text: str, caret_index: int) -> int | None:
    if caret_index >= len(text) or text[caret_index] != "^":
        return None

    value_start = caret_index + 1
    if value_start >= len(text):
        return None
    if text[value_start] == "{":
        group_end = _find_latex_group_end(text, value_start)
        return None if group_end is None else group_end + 1
    if text[value_start].isalnum():
        return value_start + 1
    return None


def _add_latex_paragraph(doc, text, style=None):
    """
    将含有 $...$ / $$...$$ 的文本插入 Word，
    普通文本 → run
    公式 → OMML
    """
    inline_math = re.compile(r'\$(.+?)\$')
    block_math = re.compile(r'\$\$(.+?)\$\$', re.DOTALL)

    # 1. 先处理块级公式 $$...$$
    pos = 0
    for m in block_math.finditer(text):
        before = text[pos:m.start()]
        if before.strip():
            p = doc.add_paragraph(before, style=style)

        latex = m.group(1).strip()
        omml = _latex_to_omml(latex)

        p = doc.add_paragraph(style=style)
        _insert_omml(p, omml)

        pos = m.end()

    # 剩余部分继续处理行内公式
    text = text[pos:]

    # 2. 行内公式处理 $...$
    p = doc.add_paragraph(style=style)
    pos = 0
    for m in inline_math.finditer(text):
        before = text[pos:m.start()]
        if before:
            p.add_run(before)

        latex = m.group(1).strip()
        omml = _latex_to_omml(latex)
        _insert_omml(p, omml)

        pos = m.end()

    # 3. 剩余普通文本
    if pos < len(text):
        p.add_run(text[pos:])


def _set_default_table_border(table):
    # 表格居中对齐
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置边框（模拟 Table Grid）
    tbl = _docx_table_element(table)
    tbl_pr = tbl.xpath('./w:tblPr')[0]

    tbl_borders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # 实线
        border.set(qn('w:sz'), '4')  # 线宽（1/8 pt）
        border.set(qn('w:space'), '0')  # 间距
        border.set(qn('w:color'), 'auto')  # 自动颜色
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)


def _add_html_table_to_doc(doc, element, context: HtmlToDocContext):
    """Add one HTML table element to a docx document."""
    table_style = _get_style_by_tag(element.name, context.style_dict, doc)
    rows = element.find_all('tr')
    if not rows:
        return

    row_cells = [row.find_all(['td', 'th']) for row in rows]
    cols_count = max((len(cells) for cells in row_cells), default=0)
    if cols_count == 0:
        return

    table = doc.add_table(rows=len(rows), cols=cols_count)
    if table_style.type == WD_STYLE_TYPE.TABLE:
        table.style = table_style
    else:
        _set_default_table_border(table)

    paragraph_style = _get_style_by_tag("p", context.style_dict, doc)
    style_r_pr = paragraph_style.element.get_or_add_rPr()
    style_r_fonts = style_r_pr.find(qn('w:rFonts'))

    for r_idx, cells in enumerate(row_cells):
        for c_idx, cell in enumerate(cells):
            target_cell = table.cell(r_idx, c_idx)
            p = target_cell.paragraphs[0]
            p.clear()
            for child in cell.contents:
                _process_inline(
                    p,
                    child,
                    replace(context, style_r_fonts=style_r_fonts),
                )


def _process_block_element(
    doc,
    element,
    context: HtmlToDocContext,
    state: HtmlBlockState = HtmlBlockState(),
):
    """Process one block-level HTML element into a docx document."""
    if element.name is None:
        return
    if state.depth >= context.max_depth:
        text = element.get_text(strip=True)
        if text:
            paragraph_style = _get_style_by_tag("p", context.style_dict, doc)
            doc.add_paragraph(text, style=paragraph_style)
        return

    if element.name in HEADING_TAGS:
        _add_para_and_apply_style(doc, element, context)

    elif element.name == 'p':
        _add_para_and_apply_style(doc, element, context)

    elif element.name == 'pre':
        paragraph_style = _get_style_by_tag("p", context.style_dict, doc)
        code_text = element.get_text()
        for line in code_text.rstrip("\n").splitlines() or [""]:
            doc.add_paragraph(line, style=paragraph_style)

    elif element.name == 'blockquote':
        para = doc.add_paragraph(element.get_text(strip=True))
        para.paragraph_format.left_indent = Pt(18)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

    elif element.name in ('ul', 'ol'):
        current_list_num_id = (
            state.list_num_id
            if state.list_num_id is not None and state.list_tag == element.name
            else _create_word_list_numbering(doc, ordered=element.name == "ol")
        )
        paragraph_style = _get_style_by_tag("p", context.style_dict, doc)
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style=paragraph_style)
            if state.list_depth:
                p.paragraph_format.left_indent = Pt(18 * state.list_depth)
            _apply_word_list_numbering(p, current_list_num_id, state.list_depth)
            style_r_pr = paragraph_style.element.get_or_add_rPr()
            style_r_fonts = style_r_pr.find(qn('w:rFonts'))
            for child in li.contents:
                if getattr(child, "name", None) in ("ul", "ol"):
                    _process_block_element(
                        doc,
                        child,
                        context,
                        replace(
                            state,
                            depth=state.depth + 1,
                            list_depth=state.list_depth + 1,
                            list_num_id=current_list_num_id,
                            list_tag=element.name,
                        ),
                    )
                    continue
                _process_inline(
                    p,
                    child,
                    replace(context, style_r_fonts=style_r_fonts),
                )

    elif element.name == 'table':
        _add_html_table_to_doc(doc, element, context)

    elif element.name in ('div', 'section', 'article', 'main'):
        for child in element.children:
            _process_block_element(
                doc,
                child,
                context,
                replace(
                    state,
                    depth=state.depth + 1,
                    list_num_id=None,
                    list_tag=None,
                ),
            )


def _get_available_page_width(doc):
    if not doc.sections:
        return None
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def html_to_doc(doc, html, style_dict, base_path: str | Path | None = None):
    """将 HTML 内容转换并写入 docx 文档对象。"""
    soup = BeautifulSoup(html, 'html.parser')
    container = soup.find("div", class_="report-container")
    if container is None:
        container = soup.body or soup

    resolved_base_path = Path(base_path).resolve() if base_path is not None else None
    max_image_width = _get_available_page_width(doc)
    context = HtmlToDocContext(
        style_dict=style_dict,
        base_path=resolved_base_path,
        max_image_width=max_image_width,
    )
    for element in container.children:
        _process_block_element(
            doc,
            element,
            context,
        )


def set_global_styles(doc, font_name="微软雅黑", font_size=11, line_spacing=1.15):
    """为 docx 文档设置全局字体与段落样式。"""
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = font_name
    normal_font.size = Pt(font_size)
    _docx_style_element(normal_style).rPr.rFonts.set(qn('w:eastAsia'), font_name)

    heading_sizes = [24, 18, 16, 14, 12, 11]
    for i in range(1, 7):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = font_name
        heading_font.size = Pt(heading_sizes[i - 1])
        heading_font.italic = False
        _docx_style_element(heading_style).rPr.rFonts.set(qn('w:eastAsia'), font_name)

    for style in doc.styles:
        if style.type == 1:  # Paragraph style
            pf = style.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = line_spacing
