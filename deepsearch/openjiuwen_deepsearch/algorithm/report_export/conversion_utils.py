# -*- coding: UTF-8 -*-
# Copyright (c) Huawei Technologies Co., Ltd. 2026-2026. All rights reserved.
"""Shared helpers for report export conversions."""

from __future__ import annotations

import base64
import html
import logging
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

import markdown

logger = logging.getLogger(__name__)

TEXT_READ_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030", "gbk")
DEFAULT_DOCX_FONT = "Microsoft YaHei"
DOCX_IMAGE_UPSCALE_FACTOR = 3
DOCX_IMAGE_CONTRAST = 1.28
DOCX_IMAGE_SHARPNESS = 1.45
DOCX_IMAGE_COLOR = 1.08
SAFE_FILENAME_RE = re.compile(r"[^\w.-]+", re.UNICODE)
NUMBERED_HEADING_RE = re.compile(
    r"^(?P<indent>\s{0,3})(?P<number>\d+(?:\.\d+)*)(?:\.\s+|\s+)(?P<title>.+?)\s*$"
)
LIST_ITEM_RE = re.compile(r"^\s{0,3}(?:[-*+]\s+|\d+\.\s+)")
INDENTED_LIST_ITEM_RE = re.compile(r"^(?P<indent>[ \t]{4,})(?P<marker>(?:[-*+]\s+|\d+\.\s+).*)$")
LIST_ITEM_WITH_INDENT_RE = re.compile(
    r"^(?P<indent>[ \t]*)(?P<marker>(?:[-*+]\s+|\d+\.\s+).*)$"
)
MARKDOWN_IMAGE_LINE_RE = re.compile(r"^[ \t]*!\[[^\]]*\]\([^)]+\)[ \t]*$")
LEGACY_FONT_CAPTION_PREFIX_RE = re.compile(
    r"^[ \t]*<font\b[^>]*\bsize\s*=\s*[\"']?2[\"']?[^>]*>",
    flags=re.IGNORECASE,
)
INDENTED_HTML_BLOCK_END_RE = re.compile(
    r"^[ \t]+</(?:div|section|article|main)>[ \t]*$",
    flags=re.IGNORECASE,
)
MARKDOWN_TABLE_ROW_RE = re.compile(r"^[ \t]{0,3}\|")
MARKDOWN_TABLE_DELIMITER_RE = re.compile(r":?-{1,}:?")
SENTENCE_END_RE = re.compile(r"[。！？?!…]$")
CITATION_RE = re.compile(r"\[\[(\d+)\]\]\((https?://[^\s)]+(?:\([^\s)]+\)[^\s)]*)*)\)")
CHECKED_CITATION_RE = re.compile(
    r"\[\s*checked_citation:\s*\d+\s*\](\[\[\d+\]\]\((?:[^()]|\([^()]*\))*\))"
)
LEGACY_CITATION_RE = re.compile(r"\[\s*citation:\s*\d+\s*\]")
REFERENCE_LINE_RE = re.compile(r"^(?P<indent>\s*)\[(\d+)\]\.\s+(.*)$", re.MULTILINE)
CENTER_CAPTION_RE = re.compile(r'<div\s+style="text-align:\s*center;?">', flags=re.IGNORECASE)
LEGACY_FONT_CAPTION_LINE_RE = re.compile(
    r'^(?P<indent>[ \t]*)<font\b[^>]*\bsize\s*=\s*["\']?2["\']?[^>]*>(?P<body>.*?)</font>'
    r'(?P<citations>(?:<sup class="citation">.*?</sup>)*)[ \t]*$',
    flags=re.IGNORECASE | re.MULTILINE,
)
EXTERNAL_LINK_RE = re.compile(
    r'<a\s+([^>]*?)href="(https?://[^"]+)"(?![^>]*\btarget=)([^>]*)>',
    flags=re.IGNORECASE,
)
CITATION_ANCHOR_RE = re.compile(
    r'(?<!<sup class="citation">)(<a\b[^>]*href="https?://[^"]+"[^>]*>\[(\d+)\]</a>)',
    flags=re.IGNORECASE,
)
HTML_TABLE_RE = re.compile(r"<table\b[^>]*>.*?</table>", flags=re.IGNORECASE | re.DOTALL)
TABLE_WRAP_OPEN_RE = re.compile(
    r'<div\b[^>]*\bclass=["\'][^"\']*\btable-wrap\b[^"\']*["\'][^>]*>\s*$',
    flags=re.IGNORECASE,
)
DOCX_TABLE_CAPTION_RE = re.compile(
    r"^(?:表\s*[\d一二三四五六七八九十]+(?:[-－—.][\d一二三四五六七八九十]+)*|"
    r"Table\s+[\w]+(?:[-－—.][\w]+)*)\s*[:：]",
    flags=re.IGNORECASE,
)
DOUBLE_ESCAPED_ENTITY_RE = re.compile(r"&amp;(#\d+;|#x[0-9a-fA-F]+;)")
# Math formula protection — NUL-based sentinels that Python-Markdown leaves untouched.
# Fenced / inline code blocks use CODE_PLACEHOLDER so they are excluded from formula matching.
MATH_PLACEHOLDER = "\x00MATH{}\x00"
MATH_PLACEHOLDER_RE = re.compile(r"\x00MATH(\d+)\x00")
CODE_PLACEHOLDER = "\x00CODE{}\x00"
CODE_PLACEHOLDER_RE = re.compile(r"\x00CODE(\d+)\x00")
FENCED_CODE_RE = re.compile(
    r"(?ms)^(?P<indent>[ \t]{0,3})(?P<fence>`{3,}|~{3,})[^\n]*\n.*?\n(?P=indent)(?P=fence)[ \t]*$"
)
INLINE_CODE_RE = re.compile(r"`+[^`\n]*`+")
BLOCK_MATH_RE = re.compile(r"\$\$.+?\$\$", re.DOTALL)
INLINE_MATH_DOLLAR_RE = re.compile(r"^\$(?!\$)(.*?)(?<!\$)\$$", re.DOTALL)
MERMAID_BLOCK_RE = re.compile(r"(?ms)^```[ \t]*mermaid[ \t]*\r?\n(.*?)\r?\n```[ \t]*$")
CHART_IMAGE_SRC_RE = re.compile(
    r"src=(?P<quote>[\"'])charts/(?P<chart_id>[A-Za-z0-9_-]+)\.png(?P=quote)"
)

try:
    from PIL import Image, ImageEnhance

    PIL_AVAILABLE = True
except Exception:  # pragma: no cover - optional dependency branch
    PIL_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_STYLE_AVAILABLE = True
except Exception:  # pragma: no cover - optional dependency branch
    DOCX_STYLE_AVAILABLE = False


@dataclass(slots=True)
class MermaidRenderStats:
    """Collect Mermaid rendering statistics for logging.

    Attributes:
        total: Mermaid 代码块总数。
        success: 成功渲染的 Mermaid 数量。
        failed: 渲染失败并回退源码块的数量。
    """

    total: int = 0
    success: int = 0
    failed: int = 0


def read_text_with_fallback(path: Path) -> str:
    """Read a text file with a short fallback encoding list.

    Args:
        path: 需要读取的文本文件路径。

    Returns:
        str: 解码后的文本内容。

    Raises:
        UnicodeDecodeError: 所有候选编码都失败时抛出。
    """
    last_error: UnicodeDecodeError | None = None
    for encoding in TEXT_READ_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc

    raise UnicodeDecodeError(
        getattr(last_error, "encoding", "unknown"),
        getattr(last_error, "object", b""),
        getattr(last_error, "start", 0),
        getattr(last_error, "end", 0),
        f"Unable to decode text file: {path}",
    )


def make_safe_filename_component(value: str, *, default: str = "document") -> str:
    """Normalize a string into a filesystem-safe filename fragment.

    Args:
        value: 原始文件名片段。
        default: 归一化后为空时使用的默认值。

    Returns:
        str: 可安全用于文件名的字符串。
    """
    cleaned = SAFE_FILENAME_RE.sub("_", value).strip("._")
    return cleaned or default


def normalize_whitespace(text: str) -> str:
    """Normalize newlines and common spacing issues in Markdown text.

    Args:
        text: 原始文本内容。

    Returns:
        str: 归一化后的文本内容。
    """
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.replace("\u00a0", " ").replace("\u3000", " ")


def replace_citations(text: str) -> str:
    """Convert citation markdown into HTML superscript links.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 处理引用格式后的文本。
    """

    def _replace(match: re.Match[str]) -> str:
        idx, url = match.group(1), match.group(2).strip()
        safe_url = html.escape(url, quote=True)
        return (
            f'<sup class="citation">'
            f'<a href="{safe_url}" target="_blank" rel="noopener noreferrer">[{idx}]</a>'
            f"</sup>"
        )

    text = CITATION_RE.sub(_replace, text)
    return re.sub(r'[ \t]+(<sup class="citation">)', r"\1", text)


def strip_internal_citation_markers(text: str) -> str:
    """移除仅供内部处理使用的 citation 控制标记。

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 清理内部标记后的 Markdown 文本。
    """
    text = CHECKED_CITATION_RE.sub(r"\1", text)
    return LEGACY_CITATION_RE.sub("", text)


def normalize_reference_lines(text: str) -> str:
    """Normalize numbered reference lines into bullet items.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 归一化后的 Markdown 文本。
    """
    lines = text.splitlines()
    result: list[str] = []
    in_reference_section = False

    for line in lines:
        stripped = line.strip().lower()
        if stripped in {
            "# 参考文献",
            "## 参考文献",
            "### 参考文献",
            "# references",
            "## references",
            "### references",
        }:
            in_reference_section = True
            result.append(line)
            continue

        if in_reference_section:
            match = REFERENCE_LINE_RE.match(line)
            if match:
                indent = match.group("indent")
                idx = match.group(2)
                content = match.group(3)
                result.append(f"{indent}- [{idx}] {content}")
                continue

            if stripped == "" or re.match(r"^\s*[-*]\s+", line):
                result.append(line)
                continue

            if re.match(r"^\s*#{1,6}\s+", line):
                in_reference_section = False

        result.append(line)

    return "\n".join(result)


def fix_center_caption_blocks(text: str) -> str:
    """Rewrite centered caption HTML blocks into markdown-aware containers.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正后的 Markdown 文本。
    """
    return CENTER_CAPTION_RE.sub('<div class="figure-caption" markdown="1">', text)


def normalize_legacy_font_caption_blocks(text: str) -> str:
    """Rewrite legacy font-based captions into block-level caption containers.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正后的 Markdown 文本。
    """

    def _replace(match: re.Match[str]) -> str:
        indent = match.group("indent")
        body = match.group("body").strip()
        citations = match.group("citations").strip()
        caption_content = f"{body}{citations}"
        return (
            f'\n{indent}<div class="figure-caption" markdown="1">\n'
            f"{indent}{caption_content}\n"
            f"{indent}</div>\n"
        )

    return LEGACY_FONT_CAPTION_LINE_RE.sub(_replace, text)


def normalize_list_boundaries(text: str) -> str:
    """Insert a blank line before list items when the source omits one.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正列表边界后的 Markdown 文本。
    """
    block_prefixes = ("#", ">", "|", "```", "<div", "<ul", "<ol", "<li")

    def _needs_blank_line_before_list(current_line: str, previous_line: str) -> bool:
        previous = previous_line.strip()
        if not LIST_ITEM_RE.match(current_line):
            return False
        if not previous:
            return False
        if LIST_ITEM_RE.match(previous_line):
            return False
        return not previous.startswith(block_prefixes)

    lines = text.split("\n")
    normalized: list[str] = []

    for line in lines:
        previous_line = normalized[-1] if normalized else ""

        if _needs_blank_line_before_list(line, previous_line):
            normalized.append("")

        normalized.append(line)

    return "\n".join(normalized)


def normalize_table_boundaries(text: str) -> str:
    """Insert a blank line before Markdown pipe tables when the source omits one.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 修正表格边界后的 Markdown 文本。
    """

    def _is_table_row(line: str) -> bool:
        return bool(MARKDOWN_TABLE_ROW_RE.match(line))

    def _is_table_delimiter(line: str) -> bool:
        if not _is_table_row(line):
            return False
        cells = [
            cell.strip().replace(" ", "")
            for cell in line.strip().strip("|").split("|")
        ]
        return len(cells) >= 2 and all(
            bool(MARKDOWN_TABLE_DELIMITER_RE.fullmatch(cell))
            for cell in cells
        )

    lines = text.split("\n")
    normalized: list[str] = []
    in_fenced_code = False

    for index, line in enumerate(lines):
        if line.lstrip().startswith("```"):
            in_fenced_code = not in_fenced_code
            normalized.append(line)
            continue

        is_table_start = (
            not in_fenced_code
            and index + 1 < len(lines)
            and _is_table_row(line)
            and _is_table_delimiter(lines[index + 1])
        )
        if is_table_start and normalized and normalized[-1].strip():
            normalized.append("")

        normalized.append(line)

    return "\n".join(normalized)


def normalize_orphan_indented_list_items(text: str) -> str:
    """Dedent report-style list items that Markdown would otherwise treat as code.

    Args:
        text: Raw Markdown text.

    Returns:
        str: Markdown text with orphan indented list items dedented.
    """
    lines = text.split("\n")
    normalized: list[str] = []
    in_fenced_code = False
    orphan_list_indent: int | None = None

    def _previous_nonempty_line() -> str:
        for previous in reversed(normalized):
            if previous.strip():
                return previous
        return ""

    def _indent_width(indent: str) -> int:
        return len(indent.expandtabs(4))

    def _line_indent_width(line: str) -> int:
        indent = re.match(r"^[ \t]*", line).group(0)
        return _indent_width(indent)

    for line in lines:
        if line.lstrip().startswith("```"):
            in_fenced_code = not in_fenced_code
            normalized.append(line)
            orphan_list_indent = None
            continue

        match = INDENTED_LIST_ITEM_RE.match(line)
        if match and not in_fenced_code:
            indent_width = _indent_width(match.group("indent"))
            if orphan_list_indent == indent_width:
                normalized.append(match.group("marker"))
                continue

            previous = _previous_nonempty_line()
            previous_indent_width = _line_indent_width(previous)
            follows_nested_content = (
                previous_indent_width > indent_width
                or (
                    previous_indent_width == indent_width
                    and INDENTED_HTML_BLOCK_END_RE.match(previous)
                )
            )
            if (
                not follows_nested_content
                and not LIST_ITEM_RE.match(previous)
                and not INDENTED_LIST_ITEM_RE.match(previous)
            ):
                orphan_list_indent = indent_width
                normalized.append(match.group("marker"))
                continue

            orphan_list_indent = None
        elif line.strip():
            orphan_list_indent = None

        normalized.append(line)

    return "\n".join(normalized)


def normalize_interrupted_nested_list_blocks(text: str) -> str:
    """Keep image or description blocks inside the list items they interrupt.

    当图片或 ``<font size=2>`` 描述块出现在嵌套列表项之间时，Markdown 会
    将其视为打断列表的独立段落，导致后续同级列表项被提升为顶级列表。
    本函数将这些打断块重新缩进到其所属父列表项内部，保证列表层级不丢失。

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 打断块已归位到父列表项内部的 Markdown 文本。
    """
    lines = text.split("\n")

    def _indent_width(indent: str) -> int:
        return len(indent.expandtabs(4))

    for index, line in enumerate(lines):
        current_item = LIST_ITEM_WITH_INDENT_RE.match(line)
        if current_item is None:
            continue

        list_indent = _indent_width(current_item.group("indent"))
        block_indexes: list[int] = []
        saw_image = False
        saw_font_description = False

        for cursor in range(index + 1, len(lines)):
            candidate = lines[cursor]
            if not candidate.strip():
                block_indexes.append(cursor)
                continue

            if MARKDOWN_IMAGE_LINE_RE.match(candidate):
                saw_image = True
                block_indexes.append(cursor)
                continue

            if saw_image and LEGACY_FONT_CAPTION_PREFIX_RE.match(candidate):
                saw_font_description = True
                block_indexes.append(cursor)
                continue

            if not saw_image and LEGACY_FONT_CAPTION_PREFIX_RE.match(candidate):
                saw_font_description = True
                block_indexes.append(cursor)
                continue

            next_item = LIST_ITEM_WITH_INDENT_RE.match(candidate)
            next_indent = _indent_width(next_item.group("indent")) if next_item else -1
            is_matching_sibling = (
                saw_image
                and next_item is not None
                and next_indent == list_indent
            )
            is_nested_after_description = (
                saw_font_description
                and next_item is not None
                and next_indent > list_indent
            )
            if is_matching_sibling or is_nested_after_description:
                content_indent_width = next_indent if is_nested_after_description else list_indent + 4
                content_indent = " " * content_indent_width
                for block_index in block_indexes:
                    if lines[block_index].strip():
                        lines[block_index] = content_indent + lines[block_index].lstrip()
            break

    return "\n".join(lines)


def render_mermaid_supplement(supplement_markdown: str) -> str:
    """Render timeline supplement markdown into an HTML helper block.

    Args:
        supplement_markdown: Mermaid 预处理阶段生成的补充说明 Markdown。

    Returns:
        str: 内嵌 HTML 容器字符串。
    """
    supplement_markdown = supplement_markdown.strip()
    if not supplement_markdown:
        return ""
    return (
        '\n<div class="timeline-notes" markdown="1">\n'
        f"{supplement_markdown}\n"
        "</div>\n"
    )


def preprocess_markdown_text(text: str) -> str:
    """Apply Markdown preprocessing before HTML or DOCX conversion.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        str: 预处理后的 Markdown 文本。
    """
    text = normalize_whitespace(text)
    text = normalize_interrupted_nested_list_blocks(text)
    text = strip_internal_citation_markers(text)
    text = replace_citations(text)
    text = normalize_reference_lines(text)
    text = fix_center_caption_blocks(text)
    text = normalize_legacy_font_caption_blocks(text)
    text = normalize_orphan_indented_list_items(text)
    text = normalize_list_boundaries(text)
    return normalize_table_boundaries(text)


def wrap_html_tables(html_text: str) -> str:
    """Wrap HTML tables with a scrollable centering container."""
    if "<table" not in html_text.lower():
        return html_text

    def _replace(match: re.Match[str]) -> str:
        prefix = html_text[max(0, match.start() - 512): match.start()]
        if TABLE_WRAP_OPEN_RE.search(prefix):
            return match.group(0)
        return f'<div class="table-wrap">{match.group(0)}</div>'

    return HTML_TABLE_RE.sub(_replace, html_text)


def _fix_double_escaped_entities(html_text: str) -> str:
    """还原被双重转义的数字 HTML 实体。

    当 Markdown 库的 md_in_html 扩展或其他处理步骤将 HTML 实体中的
    ``&`` 转义为 ``&amp;`` 时，``&#92;`` 会变成 ``&amp;#92;``，浏览器
    会将其渲染为字面文本 ``&#92;`` 而非反斜杠 ``\\``。本函数将此类
    双重转义的数字实体还原回正常形式。

    Args:
        html_text: 可能包含双重转义实体的 HTML 文本。

    Returns:
        str: 双重转义实体已还原的 HTML 文本。
    """
    return DOUBLE_ESCAPED_ENTITY_RE.sub(r"&\1", html_text)


def postprocess_html(html_text: str) -> str:
    """Postprocess generated HTML for external links, citations and table wrappers.

    Args:
        html_text: Markdown 转换后的 HTML 文本。

    Returns:
        str: 处理后的 HTML 文本。
    """

    def _replace(match: re.Match[str]) -> str:
        before = match.group(1).rstrip()
        href = re.sub(r"\s+", "", match.group(2))
        after = match.group(3).rstrip()
        attrs = " ".join(part for part in [before, f'href="{href}"', after] if part)
        return f'<a {attrs} target="_blank" rel="noopener noreferrer">'

    def _wrap_citation(match: re.Match[str]) -> str:
        return f'<sup class="citation">{match.group(1)}</sup>'

    html_text = _fix_double_escaped_entities(html_text)
    html_text = EXTERNAL_LINK_RE.sub(_replace, html_text)
    html_text = CITATION_ANCHOR_RE.sub(_wrap_citation, html_text)
    html_text = re.sub(r'[ \t]+(<sup class="citation">)', r"\1", html_text)
    html_text = re.sub(r'(</sup>)[ \t]+(<sup class="citation">)', r"\1\2", html_text)
    return wrap_html_tables(html_text)


# Simple regex to detect currency-like dollar prefixes:
# $ followed by digits (optionally comma-separated) and optional decimal.
_CURRENCY_START_RE = re.compile(r"\$(?:\d{1,3}(?:,\d{3})*|\d+)(?:\.\d+)?")


def render_markdown_html_fragment(markdown_text: str) -> str:
    """将预处理完成的 Markdown 转换为公共 HTML fragment。

    Args:
        markdown_text: 已完成格式专属 Mermaid 替换等预处理的 Markdown。

    Returns:
        包含公式、引用和表格后处理结果的 HTML fragment。
    """
    markdown_text, math_spans = protect_math_spans(markdown_text)
    html_body = markdown.markdown(
        markdown_text,
        extensions=["extra", "toc", "md_in_html"],
        output_format="html5",
    )
    return postprocess_html(restore_math_spans(html_body, math_spans))


def protect_math_spans(text: str) -> tuple[str, list[str]]:
    """Replace LaTeX math spans with placeholders before Markdown conversion.

    Python-Markdown otherwise treats characters inside ``$...$`` / ``$$...$$``
    (notably ``_`` and ``*``) as emphasis markup and injects ``<em>``/``<strong>``
    tags into the formula. Each span is swapped for a NUL-wrapped sentinel that
    Markdown leaves untouched, then restored verbatim by :func:`restore_math_spans`
    after conversion.

    Code blocks (fenced and inline) are excluded from formula matching.

    Args:
        text: 原始 Markdown 文本。

    Returns:
        tuple[str, list[str]]: (替换占位符后的文本, 按出现顺序保存的公式原文列表)。
    """
    formulas: list[str] = []
    code_spans: list[str] = []

    def _store_code(match: re.Match[str]) -> str:
        code_spans.append(match.group(0))
        return CODE_PLACEHOLDER.format(len(code_spans) - 1)

    def _store_formula(match: re.Match[str]) -> str:
        formulas.append(match.group(0))
        return MATH_PLACEHOLDER.format(len(formulas) - 1)

    # 1. Protect code blocks first (so $ inside code is never mistaken for math)
    text = FENCED_CODE_RE.sub(_store_code, text)
    text = INLINE_CODE_RE.sub(_store_code, text)

    # 2. Protect block math $$...$$ (DOTALL so multi-line blocks match)
    text = BLOCK_MATH_RE.sub(_store_formula, text)

    # 3. Protect inline math $...$
    text = _protect_inline_math_spans(text, formulas)

    # 4. Restore code blocks (their $ signs go back as-is)
    def _restore_code(match: re.Match[str]) -> str:
        index = int(match.group(1))
        if 0 <= index < len(code_spans):
            return code_spans[index]
        return match.group(0)

    text = CODE_PLACEHOLDER_RE.sub(_restore_code, text)
    return text, formulas


# Characters that signal the start of math content after a currency-like prefix.
# 包含 < > × ÷ 等比较与乘除运算符，确保 $4 < x$、$4 × 5$ 不被误判为货币。
_MATH_CONTINUATION_RE = re.compile(r"[\\_^([{=%<>×÷]")
# Math operators / commands that distinguish a formula from plain currency.
# 与 _MATH_CONTINUATION_RE 保持同步，覆盖空格后的数学特征（如 $4 < x$）。
_MATH_FEATURE_AFTER_CURRENCY_RE = re.compile(r"[+\-*/!%%^_{}=\\<>×÷]")


def _is_currency_start(text: str, start: int) -> bool:
    """Check if a ``$`` at ``start`` is a currency amount rather than inline math.

    Returns True when the ``$`` is followed by a digits-only amount with no
    closing ``$``, or when the characters after the amount don't look like
    math operators / commands.
    """
    match = re.match(r"\$(?:\d{1,3}(?:,\d{3})*|\d+)(?:\.\d+)?", text[start:])
    if not match:
        return False
    after = start + len(match.group(0))

    # Percentage → math (e.g. $79.29%, $79.29\%)
    if after < len(text) and text[after] == "%":
        return False
    if after + 1 < len(text) and text[after] == "\\" and text[after + 1] == "%":
        return False

    # Followed by math continuation → math (e.g. $1/2$, $4*3$, $5!)
    if after < len(text) and _MATH_CONTINUATION_RE.match(text[after:]):
        return False

    # If there's a nearby closing $, check if the content between has math ops
    next_dollar = text.find("$", after)
    if next_dollar != -1 and next_dollar - start <= 50:
        between = text[after:next_dollar]
        if _MATH_FEATURE_AFTER_CURRENCY_RE.search(between):
            return False

    return True


def _protect_inline_math_spans(text: str, formulas: list[str]) -> str:
    """Find and protect inline $...$ spans, skipping currency-like patterns."""
    parts: list[str] = []
    cursor = 0
    index = 0

    while index < len(text):
        start = text.find("$", index)
        if start == -1:
            break

        # Skip escaped \$ and double $$
        if _is_escaped(text, start) or _is_double_dollar(text, start):
            index = start + 1
            continue

        # Skip currency-like dollar prefixes: $5, $10, $1,200.50, etc.
        # (but keep when followed by math operators: $1/2$, $5!, $79.29%, etc.)
        if _is_currency_start(text, start):
            index = start + 1
            continue

        # Opening $ must be followed by non-whitespace, non-punctuation
        if start + 1 >= len(text) or text[start + 1].isspace():
            index = start + 1
            continue
        if text[start + 1] in "，。、；：,.!?;:)]}）】」』":
            index = start + 1
            continue

        # Find the very next unescaped, non-double $ as the closing delimiter
        end = _find_inline_math_end(text, start + 1)
        if end is None or end <= start:
            index = start + 1
            continue

        formula = text[start:end + 1]
        content = formula[1:-1].strip()

        if content and _is_likely_inline_math(content):
            parts.append(text[cursor:start])
            formulas.append(formula)
            parts.append(MATH_PLACEHOLDER.format(len(formulas) - 1))
            cursor = end + 1
            index = end + 1
        else:
            index = end + 1

    if not parts:
        return text

    parts.append(text[cursor:])
    return "".join(parts)


def _find_inline_math_end(text: str, start_index: int) -> int | None:
    """Find the closing $ for an inline math span.

    Returns the position of the very next unescaped, non-double $ that is
    preceded by non-space content (i.e., the formula content is non-empty).
    """
    index = start_index
    while index < len(text):
        end = text.find("$", index)
        if end == -1:
            return None
        # Inline math cannot span multiple lines
        if "\n" in text[start_index:end]:
            return None
        if _is_escaped(text, end) or _is_double_dollar(text, end):
            index = end + 1
            continue
        # Closing $ must be preceded by content (not whitespace)
        if end <= start_index:
            return None
        if text[end - 1].isspace():
            index = end + 1
            continue
        return end
    return None


def _is_likely_inline_math(content: str) -> bool:
    """Return whether content looks like LaTeX math rather than plain text/currency."""
    if not content:
        return False
    if len(content) > 250:
        return False

    # ---------- fast path: contains unambiguous math characters ----------
    # These characters almost never appear in plain text / currency contexts.
    math_char_re = re.compile(
        r"[\\_^{}*!×÷∑∫√≈≠≤≥<>±'|%=()\[\]]"
    )
    if math_char_re.search(content):
        return True

    # Greek letter
    if re.search(r"[Α-ω]", content):
        return True

    # ---------- alphanumeric tokens (letter + optional digits) ----------
    # Handles: G, H0, xyz, n, f, velocityvelocityvelocity, etc.
    # 上方 len(content) > 250 已做总体长度限制，此处不再对单 token 施加额外上限，
    # 否则会将 velocityvelocityvelocity 等长变量名误判为非公式（commit 867d0e6 回归）。
    if re.match(r"^[A-Za-zΑ-ω][A-Za-zΑ-ω0-9]*$", content):
        return True

    # ---------- numeric expressions ----------
    # Must contain an operator character (not just plain digits/commas).
    # Avoids matching currency values like "1,200" or "1,200.50".
    if re.match(r"^[\d+\-*/!.,\s]+$", content) and re.search(r"[-+*/!]", content):
        return True

    # Plain formatted number WITH a percentage sign
    if re.match(r"^\d+(?:[.,]\d+)?%$", content):
        return True

    # ---------- function-like: name followed by ( or [ ----------
    # Handles: f(t), P(1,2,3), \sin(x), etc.
    # Allow primes between name and parens: f'(t), f''(x)
    if re.search(
        r"^[A-Za-z\\]+(?:'++)?(?:_\{?\w+\}?|\^\{?\w+\}?)*(?:\s*,\s*[A-Za-z\\]+)*"
        r"\s*[(\[]",
        content,
    ):
        return True

    # ---------- variable list: a, b, c, ... ----------
    if re.match(r"^[A-Za-z](?:\s*,\s*[A-Za-z])+(?:\s*,\s*\.{3})?$", content):
        return True

    # ---------- parenthesized tuple: (x, y, z) ----------
    if re.match(r"^\(\s*[A-Za-z](?:\s*,\s*[A-Za-z])+\s*\)$", content):
        return True

    # ---------- equation-like: contains = with variables ----------
    # e.g. G = (V, E), A = B + C
    if re.search(r"[A-Za-zΑ-ω]\s*=", content) and re.search(r"=\s*[A-Za-zΑ-ω(\\]", content):
        return True

    return False


def _is_double_dollar(text: str, index: int) -> bool:
    """Check if $ at index is part of a $$ pair."""
    return (
        (index > 0 and text[index - 1] == "$")
        or (index + 1 < len(text) and text[index + 1] == "$")
    )


def _is_escaped(text: str, index: int) -> bool:
    """Check if character at index is preceded by an odd number of backslashes."""
    slash_count = 0
    cursor = index - 1
    while cursor >= 0 and text[cursor] == "\\":
        slash_count += 1
        cursor -= 1
    return slash_count % 2 == 1


def restore_math_spans(html_text: str, formulas: list[str]) -> str:
    """Restore math placeholders produced by :func:`protect_math_spans`.

    The formula content is HTML-escaped so that characters like ``<`` and ``>``
    inside LaTeX (e.g. ``< 1``) do not break HTML parsing. KaTeX auto-render
    reads the decoded DOM text nodes, so ``&lt;``/``&gt;`` entities are correctly
    interpreted.

    Args:
        html_text: 经 Markdown 转换、仍含公式占位符的 HTML 文本。
        formulas: :func:`protect_math_spans` 返回的公式原文列表。

    Returns:
        str: 占位符还原为公式原文后的 HTML 文本。
    """
    def _restore(match: re.Match[str]) -> str:
        index = int(match.group(1))
        if 0 <= index < len(formulas):
            return html.escape(formulas[index], quote=False)
        return match.group(0)

    return MATH_PLACEHOLDER_RE.sub(_restore, html_text)


def _neighbor_numbered_line(lines: list[str], index: int, *, reverse: bool) -> str | None:
    """Find the nearest non-empty neighboring line.

    Args:
        lines: 文本行列表。
        index: 当前行索引。
        reverse: 是否向前查找。

    Returns:
        str | None: 相邻非空行文本。
    """
    step = -1 if reverse else 1
    cursor = index + step
    while 0 <= cursor < len(lines):
        stripped = lines[cursor].strip()
        if stripped:
            return stripped
        cursor += step
    return None


def _should_promote_numbered_heading(lines: list[str], index: int, match: re.Match[str]) -> bool:
    """Decide whether a numbered line should become a Markdown heading.

    Args:
        lines: 文本行列表。
        index: 当前行索引。
        match: 编号标题匹配结果。

    Returns:
        bool: 是否应提升为标题。
    """
    title = match.group("title").strip()
    if len(title) > 80 or SENTENCE_END_RE.search(title):
        return False
    if index > 0 and lines[index - 1].strip():
        return False
    if index + 1 < len(lines) and lines[index + 1].strip():
        return False

    prev_nonempty = _neighbor_numbered_line(lines, index, reverse=True)
    next_nonempty = _neighbor_numbered_line(lines, index, reverse=False)
    if prev_nonempty and NUMBERED_HEADING_RE.match(prev_nonempty):
        return False
    if next_nonempty and NUMBERED_HEADING_RE.match(next_nonempty):
        return False
    return True


def normalize_headings(content: str) -> str:
    """Normalize numbered headings into Markdown heading syntax.

    Args:
        content: 原始 Markdown 文本。

    Returns:
        str: 标题归一化后的 Markdown 文本。
    """
    content = normalize_whitespace(content)
    lines = content.split("\n")
    out: list[str] = []
    in_code_block = False

    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            out.append(line)
            continue
        if in_code_block:
            out.append(line)
            continue
        if not stripped:
            out.append("")
            continue

        match_hash = re.match(r"^(#{1,6})\s*(.+?)\s*$", stripped)
        if match_hash:
            hashes = match_hash.group(1)
            title = match_hash.group(2).strip()
            if out and out[-1] != "":
                out.append("")
            out.append(f"{hashes} {title}")
            out.append("")
            continue

        match_numbered = NUMBERED_HEADING_RE.match(line)
        if match_numbered and _should_promote_numbered_heading(lines, index, match_numbered):
            numbering = match_numbered.group("number")
            title = match_numbered.group("title").strip()
            level = min(numbering.count(".") + 1, 6)
            heading = "#" * level + " " + f"{numbering} {title}"
            if out and out[-1] != "":
                out.append("")
            out.append(heading)
            out.append("")
            continue

        out.append(line)

    result = "\n".join(out)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip() + "\n"


def _set_rfonts(rpr, font_name: str) -> None:
    """Set all Word run font slots to a target font.

    Args:
        rpr: Word run properties element。
        font_name: 目标字体名。

    Returns:
        None.
    """
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def _set_run_font(run, font_name: str) -> None:
    """Apply a font to one docx run.

    Args:
        run: python-docx run 对象。
        font_name: 目标字体。

    Returns:
        None.
    """
    run.font.name = font_name
    _set_rfonts(run.element.get_or_add_rPr(), font_name)


def _set_style_font(style, font_name: str) -> None:
    """Apply a font to one docx style.

    Args:
        style: python-docx style 对象。
        font_name: 目标字体。

    Returns:
        None.
    """
    style.font.name = font_name
    _set_rfonts(style.element.get_or_add_rPr(), font_name)


def _apply_font_to_table(table, font_name: str) -> None:
    """Apply a font recursively to a docx table.

    Args:
        table: python-docx table 对象。
        font_name: 目标字体。

    Returns:
        None.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, font_name)
            for nested_table in cell.tables:
                _apply_font_to_table(nested_table, font_name)


def _center_docx_table(table) -> None:
    """Center one docx table and its nested tables.

    Args:
        table: python-docx table 对象。

    Returns:
        None.
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for nested_table in cell.tables:
                _center_docx_table(nested_table)


def _center_docx_table_captions(paragraphs) -> None:
    """Center paragraphs that look like table captions.

    Args:
        paragraphs: python-docx paragraph iterable.

    Returns:
        None.
    """
    for paragraph in paragraphs:
        if DOCX_TABLE_CAPTION_RE.match(paragraph.text.strip()):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def normalize_docx_tables(docx_path: Path) -> None:
    """Center tables and table-caption paragraphs in a generated DOCX file.

    Args:
        docx_path: DOCX 文件路径。

    Returns:
        None.
    """
    if not DOCX_STYLE_AVAILABLE:
        logger.warning("python-docx is unavailable, skipping table normalization for %s.", docx_path)
        return

    document = Document(docx_path)
    _center_docx_table_captions(document.paragraphs)
    for table in document.tables:
        _center_docx_table(table)

    for section in document.sections:
        _center_docx_table_captions(section.header.paragraphs)
        _center_docx_table_captions(section.footer.paragraphs)
        for table in section.header.tables:
            _center_docx_table(table)
        for table in section.footer.tables:
            _center_docx_table(table)

    document.save(docx_path)


def normalize_docx_fonts(docx_path: Path, *, font_name: str = DEFAULT_DOCX_FONT) -> None:
    """Normalize the font family across a generated DOCX file.

    Args:
        docx_path: DOCX 文件路径。
        font_name: 目标字体名。

    Returns:
        None.
    """
    if not DOCX_STYLE_AVAILABLE:
        logger.warning("python-docx is unavailable, skipping font normalization for %s.", docx_path)
        return

    document = Document(docx_path)
    for style_name in (
        "Normal",
        "Title",
        "Subtitle",
        "Body Text",
        "Hyperlink",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "Heading 7",
        "Heading 8",
        "Heading 9",
    ):
        try:
            _set_style_font(document.styles[style_name], font_name)
        except KeyError:
            continue

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font_name)

    for table in document.tables:
        _apply_font_to_table(table, font_name)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for table in section.header.tables:
            _apply_font_to_table(table, font_name)
        for table in section.footer.tables:
            _apply_font_to_table(table, font_name)

    document.save(docx_path)


def inline_chart_images(html_text: str, base_path: Path) -> str:
    """将 bundle 内的 VLM PNG 引用替换为 Data URI。

    仅处理 bundle 生成的安全 `charts/<id>.png` 路径，其他本地或远程
    图片引用保持原样。

    Args:
        html_text: 已生成的 HTML 文本。
        base_path: `report.md` 所在的 bundle 根目录。

    Returns:
        VLM PNG 已内嵌的 HTML 文本；资源不存在时保留原引用。
    """

    def _replace(match: re.Match[str]) -> str:
        chart_path = base_path / "charts" / f"{match.group('chart_id')}.png"
        if not chart_path.is_file():
            return match.group(0)
        encoded = base64.b64encode(chart_path.read_bytes()).decode("ascii")
        quote = match.group("quote")
        return f"src={quote}data:image/png;base64,{encoded}{quote}"

    return CHART_IMAGE_SRC_RE.sub(_replace, html_text)


def enhance_image(image_path: str) -> None:
    """Upscale and sharpen Mermaid PNG images for DOCX output.

    Args:
        image_path: PNG 图片路径。

    Returns:
        None.
    """
    if not PIL_AVAILABLE:
        return

    try:
        with Image.open(image_path) as original:
            if original.mode in ("RGBA", "LA"):
                background = Image.new("RGBA", original.size, (255, 255, 255, 255))
                background.alpha_composite(original.convert("RGBA"))
                image = background.convert("RGB")
            else:
                image = original.convert("RGB")

        image = image.resize(
            (image.width * DOCX_IMAGE_UPSCALE_FACTOR, image.height * DOCX_IMAGE_UPSCALE_FACTOR),
            Image.Resampling.LANCZOS,
        )
        image = ImageEnhance.Contrast(image).enhance(DOCX_IMAGE_CONTRAST)
        image = ImageEnhance.Sharpness(image).enhance(DOCX_IMAGE_SHARPNESS)
        image = ImageEnhance.Color(image).enhance(DOCX_IMAGE_COLOR)
        image.save(image_path, format="PNG", optimize=True)
    except Exception as exc:  # pragma: no cover - best effort enhancement
        logger.warning("Failed to enhance image %s: %s", image_path, exc)
